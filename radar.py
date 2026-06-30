import streamlit as st
import time
import calendar
import io
import requests
from docx import Document
import utils

def run_module():
    st.title("⚡ PickPivot: Radar Orzecznictwa")
    st.markdown("Wyszukuje interpretacje podatkowe po słowach kluczowych.")

    konfiguracja = utils.wczytaj_historie(utils.PLIK_KONFIGURACJI_M1)
    przetworzone_id = set(konfiguracja.get("przetworzone_id", []))
    ukonczone_kombinacje = set(konfiguracja.get("ukonczone_kombinacje", []))
    pelne_tresci = utils.wczytaj_pelne_tresci(utils.PLIK_REKORDOW_M1)

    if pelne_tresci:
        st.success(f"💾 BAZA RADARU: Zabezpieczono {len(pelne_tresci)} orzeczeń.")
        colA, colB = st.columns(2)
        with colA:
            if st.button("📄 GENERUJ RAPORT WORD", use_container_width=True, type="primary"):
                doc = Document()
                for r in pelne_tresci:
                    doc.add_heading(f"Sygnatura: {r['Sygnatura']}", 1)
                    doc.add_paragraph(f"Data: {r['Data']} | Podatek: {r['Podatek']}")
                    doc.add_paragraph(f"Link: {r['Link']}")
                    doc.add_paragraph(utils.wyczysc_tekst_dla_worda(r['Tekst']))
                    doc.add_page_break()
                out = io.BytesIO()
                doc.save(out)
                st.download_button("📥 Pobierz plik", out.getvalue(), "Radar.docx", use_container_width=True)
        with colB:
            if st.button("🗑️ Resetuj bazę Radaru", use_container_width=True):
                utils.wyczysc_dane_serwera(utils.PLIK_KONFIGURACJI_M1, utils.PLIK_REKORDOW_M1)
                st.rerun()
        st.markdown("---")

    col1, col2, col3 = st.columns(3)
    with col1: wybrane_lata = st.multiselect("Wybierz lata:", [2024, 2025, 2026])
    with col2: wybrane_miesiace_ui = st.multiselect("Wybierz miesiące:", utils.MIESIACE_PL)
    with col3: wybrane_podatki_ui = st.multiselect("Podatki:", ["PIT", "CIT", "VAT", "AKCYZA"], default=["VAT"])

    if st.button("🚀 Uruchom skanowanie słów", use_container_width=True):
        if not wybrane_lata or not wybrane_miesiace_ui or not wybrane_podatki_ui:
            st.error("Proszę wybrać parametry."); st.stop()
            
        wybrane_miesiace = [utils.MIESIACE_PL.index(m) + 1 for m in wybrane_miesiace_ui]
        pasek_postepu, status_tekst = st.progress(0), st.empty()
        zapytania, trafienia = 0, 0
        max_z = len(wybrane_lata) * len(wybrane_miesiace) * len(utils.FRAZY_KLUCZOWE) * len(wybrane_podatki_ui)

        with requests.Session() as sesja:
            for rok in wybrane_lata:
                for miesiac in wybrane_miesiace:
                    _, ost_dzien = calendar.monthrange(rok, miesiac)
                    ds, dk = f"{rok}-{miesiac:02d}-01", f"{rok}-{miesiac:02d}-{ost_dzien:02d}"
                    for fraza in utils.FRAZY_KLUCZOWE:
                        for podatek in wybrane_podatki_ui:
                            klucz = f"M1_{rok}_{miesiac}_{fraza}_{podatek}"
                            if klucz in ukonczone_kombinacje:
                                zapytania += 1; continue
                            
                            status_tekst.info(f"Odpytuje: {fraza} ({podatek}) dla {miesiac:02d}/{rok}...")
                            lista_trafien, _ = utils.szukaj_w_api_mf(ds, dk, fraza, sesja, podatek, utils.KODY_PRZEPISOW[podatek])
                            
                            if lista_trafien:
                                cache = utils.wczytaj_pelne_tresci(utils.PLIK_REKORDOW_M1)
                                for dok in lista_trafien:
                                    if dok["id"] not in przetworzone_id:
                                        tekst, _ = utils.pobierz_tekst_pdf(dok["id"])
                                        if tekst:
                                            cache.append({"Data": dok["data"], "Podatek": dok["typ"], "Sygnatura": dok["sygnatura"], "Link": utils.PODGLAD_URL.format(id=dok["id"]), "Tekst": tekst})
                                            przetworzone_id.add(dok["id"]); konfiguracja["przetworzone_id"].append(dok["id"])
                                            trafienia += 1
                                utils.zapisz_pelne_tresci(utils.PLIK_REKORDOW_M1, cache)
                            
                            ukonczone_kombinacje.add(klucz); konfiguracja["ukonczone_kombinacje"].append(klucz)
                            utils.zapisz_historie(utils.PLIK_KONFIGURACJI_M1, konfiguracja)
                            zapytania += 1; pasek_postepu.progress(min(1.0, zapytania / max_z))
        status_tekst.success(f"Zakończono! Zebrano {trafienia} dokumentów."); time.sleep(3); st.rerun()
