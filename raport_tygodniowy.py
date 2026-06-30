#!/usr/bin/env python3
"""
raport_tygodniowy.py — Skrypt cotygodniowy uruchamiany przez GitHub Actions.

Co robi:
  1. Wyznacza zakres poprzedniego tygodnia (poniedzialek-piatek).
  2. Dla kazdego podatku (PIT, CIT, VAT, AKCYZA):
     a) Sprawdza w Supabase czy juz mamy wszystkie interpretacje z tego okresu
        (jesli nie - pobiera brakujace z API MF i zapisuje do archiwum).
     b) Generuje plik Word ze wszystkich interpretacji danego podatku z tygodnia.
     c) Zapisuje gotowy plik do tabeli raporty_tygodniowe w Supabase.
  3. Wysyla e-mail z podsumowaniem (liczba dokumentow per podatek).

Wymagane zmienne srodowiskowe (ustawiane jako GitHub Secrets):
  SUPABASE_HOST, SUPABASE_PORT, SUPABASE_DB, SUPABASE_USER, SUPABASE_PASSWORD
  GMAIL_ADRES, GMAIL_HASLO_APLIKACJI, EMAIL_ODBIORCA

Uruchomienie lokalne (test):
  python raport_tygodniowy.py --test-tydzien 2026-01-19  (dowolna data z tygodnia ktory chcesz przetworzyc)
"""

import os
import sys
import io
import argparse
import calendar
import smtplib
import requests
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from docx import Document

import db_core
import utils  # ten sam plik co w aplikacji Streamlit - logika pobierania z MF


PODATKI = ["PIT", "CIT", "VAT", "AKCYZA"]


# ---------------------------------------------------------------------------
# KONFIGURACJA Z ZMIENNYCH SRODOWISKOWYCH
# ---------------------------------------------------------------------------
def _wczytaj_config_supabase() -> dict:
    """Czyta dane polaczenia z env vars (ustawione jako GitHub Secrets)."""
    url = os.environ.get("SUPABASE_URL", "")
    if url:
        return {"url": url}
    return {
        "host":     os.environ["SUPABASE_HOST"],
        "port":     os.environ.get("SUPABASE_PORT", "5432"),
        "database": os.environ.get("SUPABASE_DB", "postgres"),
        "user":     os.environ["SUPABASE_USER"],
        "password": os.environ["SUPABASE_PASSWORD"],
    }


# ---------------------------------------------------------------------------
# GENEROWANIE WORD (identyczne z downloader.py - spojnosc formatu)
# ---------------------------------------------------------------------------
def _generuj_word(rekordy: list, podatek: str, data_od: str, data_do: str) -> bytes:
    doc = Document()
    doc.add_heading(f"PickPivot — Raport Tygodniowy: {podatek}", 0)
    doc.add_paragraph(f"Okres: {data_od} — {data_do} (poniedzialek-piatek)")
    doc.add_paragraph(f"Wygenerowano automatycznie: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Liczba dokumentow: {len(rekordy)}")
    doc.add_page_break()

    for r in sorted(rekordy, key=lambda x: x.get("Data", ""), reverse=True):
        doc.add_heading(f"Sygnatura: {r['Sygnatura']}", 1)
        doc.add_paragraph(f"Data:    {r['Data']}")
        doc.add_paragraph(f"Podatek: {r['Podatek']}")
        doc.add_paragraph(f"Link:    {r['Link']}")
        if r.get("Format"):
            doc.add_paragraph(f"Zrodlo:  {r['Format']}")
        doc.add_paragraph(utils.wyczysc_tekst_dla_worda(r["Tekst"]))
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# POBIERANIE BRAKUJACYCH DOKUMENTOW Z MF DLA DANEGO TYGODNIA
# ---------------------------------------------------------------------------
def _uzupelnij_archiwum_dla_tygodnia(db: db_core.SupabaseDB, podatek: str,
                                       data_od: datetime, data_do: datetime) -> int:
    """
    Sprawdza czy mamy wszystkie dokumenty z danego tygodnia w archiwum.
    Jesli nie - pobiera brakujace z API MF. Zwraca liczbe nowo pobranych.
    """
    print(f"  [{podatek}] Sprawdzam API MF dla okresu {data_od.date()} — {data_do.date()}...")

    znane_id = db_core.pobierz_id_z_archiwum(db)

    with requests.Session() as sesja:
        lista, status = utils.pobierz_wszystko_z_okresu(
            data_od.strftime("%Y-%m-%d"),
            data_do.strftime("%Y-%m-%d"),
            sesja, podatek, utils.KODY_PODATKOW[podatek],
        )

    if status not in ("OK",):
        print(f"  [{podatek}] OSTRZEZENIE: API MF zwrocilo status {status}")

    do_pobrania = [d for d in lista if d["id"] not in znane_id]
    print(f"  [{podatek}] Znaleziono {len(lista)} w MF, do pobrania: {len(do_pobrania)}")

    if not do_pobrania:
        print(f"  [{podatek}] Brak nowych dokumentow - archiwum juz aktualne dla tego okresu.")
        return 0

    # Pobieranie rownolegle (ta sama funkcja co w aplikacji Streamlit)
    def on_postep(completed, total, sygnatura, status):
        if completed % 10 == 0 or completed == total:
            print(f"  [{podatek}] Postep: {completed}/{total}")

    nowe_tresci, nowe_przetworzone, nowe_uszkodzone, blokada = \
        utils.pobierz_dokumenty_rownolegle(
            do_pobrania, znane_id, set(),
            callback_postep=on_postep, workers=4,
        )

    if nowe_tresci:
        zapisanych = db_core.zapisz_wiele_do_archiwum(db, nowe_tresci, "github_actions_cron")
        print(f"  [{podatek}] Zapisano {zapisanych} nowych dokumentow do archiwum.")
        return zapisanych

    if blokada:
        print(f"  [{podatek}] UWAGA: Wykryto blokade IP podczas pobierania.")

    return 0


# ---------------------------------------------------------------------------
# WYSYLKA EMAIL
# ---------------------------------------------------------------------------
def _wyslij_email_podsumowanie(wyniki: dict, data_od: datetime, data_do: datetime,
                                  sukces: bool, nowych_w_tej_probie: int = 0):
    """
    wyniki: {"VAT": 45, "PIT": 12, "CIT": 8, "AKCYZA": 0}
    nowych_w_tej_probie: ile dokumentow przybylo w TEJ konkretnej probie (nie lacznie w archiwum)
    """
    gmail_adres = os.environ.get("GMAIL_ADRES")
    gmail_haslo = os.environ.get("GMAIL_HASLO_APLIKACJI")
    odbiorca    = os.environ.get("EMAIL_ODBIORCA", gmail_adres)

    if not gmail_adres or not gmail_haslo:
        print("Brak konfiguracji email (GMAIL_ADRES / GMAIL_HASLO_APLIKACJI) — pomijam wysylke.")
        return

    tytul_status = "✅ Sukces" if sukces else "⚠️ Zakonczono z ostrzezeniami"
    temat = f"PickPivot — Raport tygodniowy {data_od.strftime('%d.%m')}–{data_do.strftime('%d.%m.%Y')} [{tytul_status}]"

    info_nowych = (
        f"<p><b>Nowych dokumentow w tej probie:</b> {nowych_w_tej_probie}</p>"
        if nowych_w_tej_probie > 0 else
        "<p style='color:#888;'>To podsumowanie kontrolne — brak nowych dokumentow od ostatniej proby.</p>"
    )

    tresc_html = f"""
    <html><body style="font-family: Arial, sans-serif;">
    <h2>📦 PickPivot — Cotygodniowy Raport Interpretacji</h2>
    <p><b>Okres:</b> {data_od.strftime('%d.%m.%Y')} (poniedzialek) — {data_do.strftime('%d.%m.%Y')} (piatek)</p>
    <p><b>Status:</b> {tytul_status}</p>
    {info_nowych}
    <h3>Stan archiwum dla tego tygodnia (lacznie):</h3>
    <table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
        <tr style="background-color: #2C3E50; color: white;">
            <th>Podatek</th><th>Liczba dokumentow</th>
        </tr>
        {"".join(f'<tr><td>{pod}</td><td>{n}</td></tr>' for pod, n in wyniki.items())}
    </table>
    <p style="margin-top: 20px;">
        Pliki Word sa gotowe do pobrania w aplikacji PickPivot, w module
        <b>"Raporty Tygodniowe"</b>.
    </p>
    <p style="color: #888; font-size: 12px;">
        Wiadomosc wygenerowana automatycznie przez GitHub Actions.
    </p>
    </body></html>
    """

    msg = MIMEMultipart("alternative")
    msg["Subject"] = temat
    msg["From"]    = gmail_adres
    msg["To"]      = odbiorca
    msg.attach(MIMEText(tresc_html, "html", "utf-8"))

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_adres, gmail_haslo)
            server.send_message(msg)
        print(f"E-mail z podsumowaniem wyslany do {odbiorca}.")
    except Exception as e:
        print(f"BLAD wysylki email: {e}")


# ---------------------------------------------------------------------------
# GLOWNA LOGIKA
# ---------------------------------------------------------------------------
def main(data_testowa: str = None):
    print("=" * 70)
    print("PickPivot — Cotygodniowy Raport Interpretacji Podatkowych")
    print("=" * 70)

    # Wyznacz zakres tygodnia
    if data_testowa:
        dzien_ref = datetime.strptime(data_testowa, "%Y-%m-%d")
        # Dla testow: traktuj podana date jako "dzisiaj" i licz tydzien w ktorym ona jest
        data_od = dzien_ref - __import__("datetime").timedelta(days=dzien_ref.weekday())
        data_do = data_od + __import__("datetime").timedelta(days=4)
        klucz   = db_core.klucz_tygodnia(data_od)
    else:
        data_od, data_do, klucz = db_core.zakres_poprzedniego_tygodnia()

    print(f"Tydzien: {klucz} | Zakres: {data_od.date()} ({data_od.strftime('%A')}) — {data_do.date()} ({data_do.strftime('%A')})")

    # Polaczenie z baza
    config = _wczytaj_config_supabase()
    db = db_core.SupabaseDB(config)
    db.inicjalizuj_schemat()
    print("Polaczenie z Supabase OK.")

    wyniki = {}
    sukces_calosc = True
    nowych_lacznie = 0

    for podatek in PODATKI:
        print(f"\n--- {podatek} ---")
        try:
            # 1. Uzupelnij archiwum o brakujace dokumenty z tego tygodnia
            nowych = _uzupelnij_archiwum_dla_tygodnia(db, podatek, data_od, data_do)
            nowych_lacznie += nowych

            # 2. Pobierz wszystkie dokumenty z archiwum dla tego tygodnia/podatku
            rekordy = db_core.pobierz_rekordy_z_archiwum(
                db, podatek=podatek,
                data_od=data_od.strftime("%Y-%m-%d"),
                data_do=data_do.strftime("%Y-%m-%d"),
            )
            wyniki[podatek] = len(rekordy)
            print(f"  [{podatek}] Lacznie w archiwum dla tego tygodnia: {len(rekordy)} dokumentow.")

            if not rekordy:
                print(f"  [{podatek}] Brak dokumentow - pomijam generowanie pliku.")
                continue

            # 3. Generuj plik Word
            plik_bytes  = _generuj_word(rekordy, podatek, data_od.strftime("%Y-%m-%d"), data_do.strftime("%Y-%m-%d"))
            nazwa_pliku = f"Raport_{podatek}_{klucz}_{data_od.strftime('%Y%m%d')}-{data_do.strftime('%Y%m%d')}.docx"

            # 4. Zapisz do Supabase
            db_core.zapisz_raport_tygodniowy(
                db, tydzien_klucz=klucz,
                data_od=data_od.strftime("%Y-%m-%d"),
                data_do=data_do.strftime("%Y-%m-%d"),
                podatek=podatek,
                liczba_dok=len(rekordy),
                plik_bytes=plik_bytes,
                nazwa_pliku=nazwa_pliku,
            )
            print(f"  [{podatek}] Zapisano raport: {nazwa_pliku} ({len(plik_bytes)/1024:.1f} KB)")

        except Exception as e:
            print(f"  [{podatek}] BLAD: {e}")
            sukces_calosc = False
            wyniki[podatek] = wyniki.get(podatek, 0)

    print("\n" + "=" * 70)
    print("PODSUMOWANIE:")
    for pod, n in wyniki.items():
        print(f"  {pod}: {n} dokumentow")
    print(f"  Nowych dokumentow w tej probie: {nowych_lacznie}")
    print("=" * 70)

    # ── INTELIGENTNE POMIJANIE E-MAILA ──────────────────────────────────
    # Przy 20 probach w weekend nie chcemy spamowac skrzynki za kazdym razem.
    # Wysylamy e-mail TYLKO gdy:
    #   - pojawily sie nowe dokumenty w tej probie (cos faktycznie sie zmienilo), LUB
    #   - wystapil blad (trzeba wiedziec), LUB
    #   - to pierwsza proba w oknie (sobota 15:00) - potwierdzenie startu, LUB
    #   - to ostatnia proba w oknie (poniedzialek ~02:00) - finalne podsumowanie
    teraz = datetime.now()
    jest_pierwsza_proba = (teraz.weekday() == 5 and teraz.hour < 16)   # sobota, wczesnie
    jest_ostatnia_proba  = (teraz.weekday() == 0 and teraz.hour < 3)    # poniedzialek, wczesnie

    wymuszony_test = data_testowa is not None  # reczne/testowe uruchomienie - zawsze wysylaj

    wyslac_email = (
        nowych_lacznie > 0
        or not sukces_calosc
        or jest_pierwsza_proba
        or jest_ostatnia_proba
        or wymuszony_test
    )

    if wyslac_email:
        _wyslij_email_podsumowanie(wyniki, data_od, data_do, sukces_calosc, nowych_lacznie)
    else:
        print("Pomijam wysylke e-mail - brak nowych dokumentow w tej probie (nie pierwsza/ostatnia proba w oknie).")

    if not sukces_calosc:
        sys.exit(1)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--test-tydzien", type=str, default=None,
        help="Data YYYY-MM-DD z dowolnego dnia tygodnia ktory chcesz przetworzyc (do testow lokalnych)"
    )
    args = parser.parse_args()
    main(data_testowa=args.test_tydzien)
