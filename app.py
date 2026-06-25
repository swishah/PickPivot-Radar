import streamlit as st
import requests
import PyPDF2
import time
import random
import io
import calendar
import json
import os
import re
import pandas as pd
import numpy as np
import yfinance as yf
import warnings
from datetime import datetime, date
from docx import Document

# IMPORT NASZEGO NOWEGO, ODSEPAROWANEGO PLIKU:
import cfo_analyzer 

warnings.filterwarnings("ignore", category=FutureWarning)
warnings.filterwarnings("ignore", category=UserWarning)

st.set_page_config(page_title="PickPivot Platform", page_icon="⚡", layout="wide")

if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False

if not st.session_state['authenticated']:
    st.markdown("<br><br>", unsafe_allow_html=True)
    col_login, _ = st.columns([1, 2])
    
    with col_login:
        st.title("🔐 Panel PickPivot")
        st.markdown("Dostęp do platformy jest szyfrowany i wymaga autoryzacji.")
        
        username = st.text_input("Login (Nazwa użytkownika):")
        password = st.text_input("Hasło:", type="password")
        
        if st.button("🚀 Zaloguj się", use_container_width=True, type="primary"):
            if username == "DORADCA" and password == "kontotestowe413":
                st.session_state['authenticated'] = True
                st.success("Autoryzacja pomyślna! Ładowanie platformy...")
                time.sleep(1)
                st.rerun()
            else:
                st.error("Wprowadzono niepoprawny login lub hasło.")
    st.stop() 

FOLDER_DOCELOWY = 'PickPivot_Data'
if not os.path.exists(FOLDER_DOCELOWY):
    os.makedirs(FOLDER_DOCELOWY)

PLIK_KONFIGURACJI_M1 = f"{FOLDER_DOCELOWY}/historia_m1.json"
PLIK_REKORDOW_M1 = f"{FOLDER_DOCELOWY}/baza_tresci_m1.json"

PLIK_KONFIGURACJI_M2 = f"{FOLDER_DOCELOWY}/historia_m2.json"
PLIK_REKORDOW_M2 = f"{FOLDER_DOCELOWY}/baza_tresci_m2.json"

SEARCH_API_URL_BASE = "https://eureka.mf.gov.pl/api/public/v1/wyszukiwarka/informacje/?size=100&page={page}&sort=parametryPozycjonowania%2Casc"
PDF_API_URL = "https://eureka.mf.gov.pl/api/public/v1/informacje/{id}/eksport/pdf"
PODGLAD_URL = "https://eureka.mf.gov.pl/informacje/podglad/{id}"

FRAZY_KLUCZOWE = [
    "sieć ciepłownicza", "przebudowa sieci", "przyłącze", "węzeł cieplny",
    "taryfa dla ciepła", "wodociąg", "kanalizacja", "oczyszczalnia ścieków",
    "stacja uzdatniania", "spółka komunalna"
]

KODY_PODATKOW = {
    "PIT": ".4011.", "CIT": ".4010.", "VAT": ".4012.", "AKCYZA": ".4013."
}

MIESIACE_PL = [
    "Styczeń", "Luty", "Marzec", "Kwiecień", "Maj", "Czerwiec",
    "Lipiec", "Sierpień", "Wrzesień", "Październik", "Listopad", "Grudzień"
]

# SŁOWNIKI DLA MODUŁU 3 (Skaner Giełdowy)
WIG20_MAP = {"ALE.WA": "Allegro", "ALR.WA": "Alior Bank", "BDX.WA": "Budimex", "BHW.WA": "Bank Handlowy", "CDR.WA": "CD Projekt", "CPS.WA": "Cyfrowy Polsat", "DNP.WA": "Dino Polska", "JSW.WA": "JSW", "KGH.WA": "KGHM", "KRU.WA": "Kruk", "LPP.WA": "LPP", "MBK.WA": "mBank", "OPL.WA": "Orange Polska", "PEO.WA": "Pekao SA", "PGE.WA": "PGE", "PKO.WA": "PKO BP", "PKN.WA": "ORLEN", "PZU.WA": "PZU", "SPL.WA": "Santander BP", "MDV.WA": "Modivo"}
MWIG40_MAP = {"11B.WA": "11 bit studios", "1AT.WA": "Atal", "ABS.WA": "Asseco BS", "APR.WA": "Auto Partner", "ASB.WA": "ASBIS", "BFT.WA": "Benefit Systems", "CAR.WA": "Inter Cars", "CIG.WA": "CI Games", "CLN.WA": "Celon Pharma", "COG.WA": "Cognor", "DAT.WA": "DataWalk", "DOM.WA": "Dom Development", "EAT.WA": "AmRest", "ENP.WA": "Enea", "EUR.WA": "Eurocash", "GPP.WA": "Grupa Pracuj", "GRN.WA": "Grenevia", "GTC.WA": "GTC", "HUU.WA": "Huuuge", "ING.WA": "ING BSK", "TXT.WA": "Text S.A.", "MIL.WA": "Millennium", "MBR.WA": "Mo-BRUK", "NEU.WA": "Neuca", "PLW.WA": "PlayWay", "RVU.WA": "Revuele", "SEL.WA": "Selena FM", "STP.WA": "Stalproduct", "TEN.WA": "Ten Square Games", "TPE.WA": "Tauron", "VRG.WA": "VRG", "WPL.WA": "Wirtualna Polska", "XTB.WA": "XTB", "GPW.WA": "GPW", "SNK.WA": "Sanok", "AST.WA": "Asseco POL", "ATC.WA": "Arctic Paper"}
DAX_MAP = {"ADS.DE": "Adidas", "AIR.DE": "Airbus", "ALV.DE": "Allianz", "BAS.DE": "BASF", "BAYN.DE": "Bayer", "BEI.DE": "Beiersdorf", "BMW.DE": "BMW", "BNR.DE": "Brenntag", "CBK.DE": "Commerzbank", "CON.DE": "Continental", "1COV.DE": "Covestro", "DTG.DE": "Daimler Truck", "DBK.DE": "Deutsche Bank", "DB1.DE": "Deutsche Börse", "DPW.DE": "DHL Group", "DTE.DE": "Deutsche Telekom", "EOAN.DE": "E.ON", "FRE.DE": "Fresenius", "HNR1.DE": "Hannover Re", "HEI.DE": "Heidelberg Materials", "HEN3.DE": "Henkel", "IFX.DE": "Infineon", "MBG.DE": "Mercedes-Benz", "MRK.DE": "Merck", "MTX.DE": "MTU Aero Engines", "MUV2.DE": "Munich Re", "P911.DE": "Porsche AG", "PAH3.DE": "Porsche SE", "QIA.DE": "Qiagen", "RHM.DE": "Rheinmetall", "RWE.DE": "RWE", "SAP.DE": "SAP", "SRT3.DE": "Sartorius", "SIE.DE": "Siemens", "ENR.DE": "Siemens Energy", "SHL.DE": "Siemens Healthineers", "SY1.DE": "Symrise", "VOW3.DE": "Volkswagen", "VNA.DE": "Vonovia", "ZAL.DE": "Zalando"}
CAC40_MAP = {"AC.PA": "Accor", "AI.PA": "Air Liquide", "AIR.PA": "Airbus", "MT.AS": "ArcelorMittal", "CS.PA": "AXA", "BNP.PA": "BNP Paribas", "EN.PA": "Bouygues", "CAP.PA": "Capgemini", "CA.PA": "Carrefour", "ACA.PA": "Crédit Agricole", "BN.PA": "Danone", "DSY.PA": "Dassault Systèmes", "EDEN.PA": "Edenred", "ENGI.PA": "Engie", "EL.PA": "EssilorLuxottica", "ERF.PA": "Eurofins Scientific", "RMS.PA": "Hermès", "KER.PA": "Kering", "LR.PA": "Legrand", "OR.PA": "L'Oréal", "MC.PA": "LVMH", "ML.PA": "Michelin", "ORP.PA": "Orange", "PRV.PA": "Pernod Ricard", "PUB.PA": "Publicis Groupe", "RNO.PA": "Renault", "SAF.PA": "Safran", "SGO.PA": "Saint-Gobain", "SAN.PA": "Sanofi", "SU.PA": "Schneider Electric", "GLE.PA": "Société Générale", "STLAP.PA": "Stellantis", "STMPA.PA": "STMicroelectronics", "TEP.PA": "Teleperformance", "HO.PA": "Thales", "TTE.PA": "TotalEnergies", "URW.AS": "Unibail-Rodamco-Westfield", "VIE.PA": "Veolia", "DG.PA": "Vinci", "VIV.PA": "Vivendi"}
FTSE_MAP = {"SHEL.L": "Shell", "AZN.L": "AstraZeneca", "HSBA.L": "HSBC", "ULVR.L": "Unilever", "BP.L": "BP", "GSK.L": "GSK", "DGE.L": "Diageo", "REL.L": "RELX", "BATS.L": "British American Tobacco", "GLEN.L": "Glencore", "RIO.L": "Rio Tinto", "BA.L": "BAE Systems", "CPG.L": "Compass Group", "LSEG.L": "LSEG", "NWG.L": "NatWest Group", "BARC.L": "Barclays", "STAN.L": "Standard Chartered", "NG.L": "National Grid", "AHT.L": "Ashtead", "TSCO.L": "Tesco", "LLOY.L": "Lloyds", "PRU.L": "Prudential", "AV.L": "Aviva", "SSE.L": "SSE", "LGEN.L": "Legal & General", "RTO.L": "Rentokil", "NXT.L": "Next", "WPP.L": "WPP", "VOD.L": "Vodafone", "RR.L": "Rolls-Royce", "EZJ.L": "easyJet", "IAG.L": "IAG"}
IBEX_MAP = {"ANA.MC": "Acciona", "ACX.MC": "Acerinox", "ACS.MC": "ACS", "AENA.MC": "Aena", "AMS.MC": "Amadeus", "BKT.MC": "Bankinter", "BBVA.MC": "BBVA", "CABK.MC": "CaixaBank", "CLNX.MC": "Cellnex", "ENG.MC": "Enagás", "ELE.MC": "Endesa", "FER.MC": "Ferrovial", "FDR.MC": "Fluidra", "GRF.MC": "Grifols", "IAG.MC": "IAG", "IBE.MC": "Iberdrola", "ITX.MC": "Inditex", "IDR.MC": "Indra", "COL.MC": "Inmobiliaria Colonial", "LOG.MC": "Logista", "MAP.MC": "Mapfre", "MEL.MC": "Meliá Hotels", "MRL.MC": "Merlin Properties", "NTGY.MC": "Naturgy", "RED.MC": "Redeia", "REP.MC": "Repsol", "ROVI.MC": "Rovi", "SAB.MC": "Sabadell", "SAN.MC": "Banco Santander", "SCYR.MC": "Sacyr", "TEF.MC": "Telefónica", "UNI.MC": "Unicaja"}
OMX_MAP = {"ABB.ST": "ABB", "ALFA.ST": "Alfa Laval", "ASSA-B.ST": "ASSA ABLOY", "ATCO-A.ST": "Atlas Copco A", "ATCO-B.ST": "Atlas Copco B", "AZN.ST": "AstraZeneca", "BOL.ST": "Boliden", "ELUX-B.ST": "Electrolux", "ERIC-B.ST": "Ericsson", "ESSITY-B.ST": "Essity", "EVO.ST": "Evolution", "GETI-B.ST": "Getinge", "HEXA-B.ST": "Hexagon", "HM-B.ST": "H&M", "INVE-B.ST": "Investor B", "KINV-B.ST": "Kinnevik", "NDA-SE.ST": "Nordea", "SAND.ST": "Sandvik", "SCA-B.ST": "SCA", "SEB-A.ST": "SEB", "SHB-A.ST": "Handelsbanken", "SKA-B.ST": "Skanska", "SKF-B.ST": "SKF", "STE-R.ST": "Stora Enso", "SWED-A.ST": "Swedbank", "SWMA.ST": "Swedish Match", "TEL2-B.ST": "Tele2", "TELIA.ST": "Telia", "VOLV-B.ST": "Volvo B"}
OBX_MAP = {"EQNR.OL": "Equinor", "DNB.OL": "DNB Bank", "AKBP.OL": "Aker BP", "TEL.OL": "Telenor", "NHY.OL": "Norsk Hydro", "MOWI.OL": "Mowi", "YAR.OL": "Yara International", "ORK.OL": "Orkla", "SUBC.OL": "Subsea 7", "TOM.OL": "Tomra Systems", "STB.OL": "Storebrand", "SALM.OL": "SalMar", "GJFS.OL": "Gjensidige", "AKER.OL": "Aker", "SCHA.OL": "Schibsted A", "FRO.OL": "Frontline", "TGS.OL": "TGS", "BAKKA.OL": "Bakkafrost", "LSG.OL": "Lerøy Seafood", "KOG.OL": "Kongsberg Gruppen", "NOD.OL": "Nordic Semiconductor", "NEL.OL": "Nel", "VAR.OL": "Vår Energi", "MPCC.OL": "MPC Container Ships"}

def wczytaj_historie(plik):
    if os.path.exists(plik):
        with open(plik, 'r', encoding='utf-8') as f:
            dane = json.load(f)
            if "uszkodzone_id" not in dane: dane["uszkodzone_id"] = []
            return dane
    return {"przetworzone_id": [], "ukonczone_kombinacje": [], "uszkodzone_id": []}

def zapisz_historie(plik, konfiguracja):
    with open(plik, 'w', encoding='utf-8') as f:
        json.dump(konfiguracja, f, ensure_ascii=False, indent=4)

def wczytaj_pelne_tresci(plik):
    if os.path.exists(plik):
        with open(plik, 'r', encoding='utf-8') as f: return json.load(f)
    return []

def zapisz_pelne_tresci(plik, lista_rekordow):
    with open(plik, 'w', encoding='utf-8') as f:
        json.dump(lista_rekordow, f, ensure_ascii=False, indent=4)

def wyczysc_dane_serwera(plik_konf, plik_rekordow):
    if os.path.exists(plik_konf): os.remove(plik_konf)
    if os.path.exists(plik_rekordow): os.remove(plik_rekordow)

def wyczysc_tekst_dla_worda(tekst):
    if not tekst: return ""
    return re.sub(r'[^\x09\x0A\x0D\x20-\x7E\x85\xA0-\uD7FF\uE000-\uFFFD\u10000-\u10FFFF]', '', tekst)

def pobierz_tekst_pdf(id_dokumentu):
    url = PDF_API_URL.format(id=id_dokumentu)
    headers_pdf = {"User-Agent": "Mozilla/5.0", "Referer": "https://eureka.mf.gov.pl/"}
    for proba in range(3):
        try:
            response = requests.get(url, headers=headers_pdf, timeout=20)
            if response.status_code == 200:
                plik_w_pamieci = io.BytesIO(response.content)
                tekst_dokumentu = ""
                reader = PyPDF2.PdfReader(plik_w_pamieci)
                for strona in reader.pages:
                    wyc = strona.extract_text()
                    if wyc: tekst_dokumentu += wyc + "\n"
                return tekst_dokumentu, "OK"
            elif response.status_code in [404, 400]: return None, "BRAK_PLIKU"
            elif response.status_code == 429: time.sleep(5)
            else: time.sleep(2)
        except:
            time.sleep(3)
    return None, "BLOKADA"

def szukaj_w_api_mf(data_start_str, data_koniec_str, fraza, sesja, nazwa_podatku, kod_sygnatury):
    dokumenty_podatkowe = []
    page = 0
    while True:
        url = SEARCH_API_URL_BASE.format(page=page)
        payload = {
            "query": fraza,
            "filter": {"KATEGORIA_INFORMACJI": [1], "DT_WYD_start": data_start_str, "DT_WYD_end": data_koniec_str},
            "columns": ["SYG", "ID_INFORMACJI", "DT_WYD"],
            "searchInFullPhrase": False, "searchInContent": True, "searchInSynonyms": True, "warunkiDodatkowe": []
        }
        headers = {"User-Agent": "Mozilla/5.0", "Content-Type": "application/json"}
        try:
            response = sesja.post(url, json=payload, headers=headers, timeout=12)
            if response.status_code == 200:
                dane = response.json()
                wyniki = dane.get('content') or dane.get('items') or []
                if not wyniki:
                    for k, v in dane.items():
                        if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                            if 'id' in v[0] or 'ID_INFORMACJI' in v[0]:
                                wyniki = v
                                break
                for d in wyniki:
                    sygnatura = str(d.get('SYG', '')).upper()
                    data_wydania = str(d.get('DT_WYD', '')).split('T')[0]
                    if kod_sygnatury in sygnatura:
                        doc_id = str(d.get('id') or d.get('ID_INFORMACJI'))
                        if doc_id: dokumenty_podatkowe.append({"id": doc_id, "sygnatura": sygnatura, "typ": nazwa_podatku, "data": data_wydania})
                if len(wyniki) < 100: break
                page += 1
                time.sleep(0.2)
            else: return dokumenty_podatkowe, "ERROR"
        except requests.exceptions.Timeout: return dokumenty_podatkowe, "TIMEOUT"
        except: return dokumenty_podatkowe, "ERROR"
    return dokumenty_podatkowe, "OK"

def pobierz_wszystko_z_okresu(data_start_str, data_koniec_str, sesja, nazwa_podatku, kod_sygnatury):
    dokumenty_podatkowe = []
    page = 0
    while True:
        url = SEARCH_API_URL_BASE.format(page=page)
        payload = {
            "filter": {"KATEGORIA_INFORMACJI": [1], "DT_WYD_start": data_start_str, "DT_WYD_end": data_koniec_str},
            "columns": ["SYG", "ID_INFORMACJI", "DT_WYD"],
            "searchInFullPhrase": False, "searchInContent": False, "searchInSynonyms": False, "warunkiDodatkowe": []
        }
        headers = {"User-Agent": "Mozilla/5.0", "Content-Type": "application/json"}
        try:
            response = sesja.post(url, json=payload, headers=headers, timeout=15)
            if response.status_code == 200:
                dane = response.json()
                wyniki = dane.get('content') or dane.get('items') or []
                if not wyniki:
                    for k, v in dane.items():
                        if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                            if 'id' in v[0] or 'ID_INFORMACJI' in v[0]:
                                wyniki = v
                                break
                for d in wyniki:
                    sygnatura = str(d.get('SYG', '')).upper()
                    data_wydania = str(d.get('DT_WYD', '')).split('T')[0]
                    if kod_sygnatury in sygnatura:
                        doc_id = str(d.get('id') or d.get('ID_INFORMACJI'))
                        if doc_id: dokumenty_podatkowe.append({"id": doc_id, "sygnatura": sygnatura, "typ": nazwa_podatku, "data": data_wydania})
                if len(wyniki) < 100: break
                page += 1
                time.sleep(0.2)
            else: return dokumenty_podatkowe, "ERROR"
        except requests.exceptions.Timeout: return dokumenty_podatkowe, "TIMEOUT"
        except: return dokumenty_podatkowe, "ERROR"
    return dokumenty_podatkowe, "OK"

@st.cache_data(ttl=3600*24)
def get_sp500_map():
    url = 'https://en.wikipedia.org/wiki/List_of_S%26P_500_companies'
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        response = requests.get(url, headers=headers)
        table = pd.read_html(io.StringIO(response.text))[0]
        table['Symbol'] = table['Symbol'].str.replace('.', '-', regex=False)
        return dict(zip(table['Symbol'], table['Security']))
    except:
        return {"AAPL": "Apple Inc.", "MSFT": "Microsoft Corp."}

def run_monte_carlo(data):
    returns = data['Close'].pct_change().dropna()
    if len(returns) < 30: return 0, 0
    mu, sigma = returns.mean(), returns.std()
    last = float(data['Close'].iloc[-1])
    sims = [last * (1 + np.random.normal(mu, sigma, 30)).prod() for _ in range(30)]
    med = np.median(sims)
    return round(med, 2), round(((med - last) / last) * 100, 1)

@st.cache_data(ttl=3600, show_spinner=False)
def analyze_market(ticker_map, label):
    results = []
    curr_y = datetime.now().year
    y1, y2, y3 = curr_y - 1, curr_y - 2, curr_y - 3

    for t, full_name in ticker_map.items():
        try:
            df = yf.download(t, period="max", interval="1d", progress=False, auto_adjust=True, actions=True)
            if df.empty or len(df) < 200: continue
            if isinstance(df.columns, pd.MultiIndex): df.columns = df.columns.get_level_values(0)

            p = float(df['Close'].iloc[-1])

            try:
                ticker_obj = yf.Ticker(t)
                info = ticker_obj.info
            except: info = {}

            def safe_get(key, is_pct=False):
                val = info.get(key)
                if val is None or pd.isna(val): return "BRAK"
                return round(val * 100, 2) if is_pct else round(val, 2)

            pe = safe_get('trailingPE')
            fwd_pe = safe_get('forwardPE')
            pb = safe_get('priceToBook')
            roe = safe_get('returnOnEquity', is_pct=True)
            op_margin = safe_get('operatingMargins', is_pct=True)
            debt_eq = safe_get('debtToEquity') 
            eps_growth = safe_get('earningsGrowth', is_pct=True)

            int_pe = "⚪ NEUTRALNIE"
            if isinstance(pe, (int, float)):
                if pe < 15: int_pe = "🟢 TANIO"
                elif pe > 25: int_pe = "🔴 DROGO"

            int_roe = "⚪ NEUTRALNIE"
            if isinstance(roe, (int, float)):
                if roe > 15: int_roe = "🟢 WYSOKIE"
                elif roe < 5: int_roe = "🔴 NISKIE"

            int_eps = "⚪ NEUTRALNIE"
            if isinstance(eps_growth, (int, float)):
                if eps_growth > 0: int_eps = "🟢 WZROST"
                elif eps_growth < 0: int_eps = "🔴 SPADEK"

            div_y1_str, div_y2_str, div_y3_str, div_yield = "BRAK", "BRAK", "BRAK", "BRAK"
            if 'Dividends' in df.columns:
                divs = df['Dividends']
                if not divs.empty and divs.sum() > 0:
                    divs_by_year = divs.groupby(divs.index.year).sum()
                    div_y1 = round(float(divs_by_year.get(y1, 0)), 2)
                    div_y2 = round(float(divs_by_year.get(y2, 0)), 2)
                    div_y3 = round(float(divs_by_year.get(y3, 0)), 2)
                    if div_y1 > 0:
                        div_y1_str = str(div_y1)
                        if p > 0: div_yield = round((div_y1 / p) * 100, 2)
                    if div_y2 > 0: div_y2_str = str(div_y2)
                    if div_y3 > 0: div_y3_str = str(div_y3)

            ath_val = df['High'].max()
            atl_val = df['Low'].min()
            ath = float(ath_val.item() if hasattr(ath_val, 'item') else ath_val)
            atl = float(atl_val.item() if hasattr(atl_val, 'item') else atl_val)
            pct_from_ath = round(((p - ath) / ath) * 100, 1)

            delta = df['Close'].diff()
            rsi = 100 - (100 / (1 + (delta.where(delta > 0, 0).rolling(14).mean() / -delta.where(delta < 0, 0).rolling(14).mean()))).iloc[-1]
            macd = (df['Close'].ewm(span=12).mean() - df['Close'].ewm(span=26).mean())
            macd_s = macd.ewm(span=9).mean()
            sma20 = float(df['Close'].rolling(20).mean().iloc[-1])
            sma50 = float(df['Close'].rolling(50).mean().iloc[-1])
            sma100 = float(df['Close'].rolling(100).mean().iloc[-1])
            sma200 = float(df['Close'].rolling(200).mean().iloc[-1])
            std20 = df['Close'].rolling(20).std().iloc[-1]
            b_low = sma20 - (std20 * 2)
            b_up = sma20 + (std20 * 2)
            b_pct = (p - b_low) / (b_up - b_low) if (b_up - b_low) != 0 else 0.5
            v_rat = float(df['Volume'].iloc[-1] / df['Volume'].rolling(20).mean().iloc[-1])
            mc_t, mc_p = run_monte_carlo(df)
            p_low = df['Low'].shift(1).rolling(20).min().iloc[-1]
            smc = "💎 SMC BUY" if (df['Low'].iloc[-1] < p_low and p > p_low and v_rat > 1.2) else "Neutralny"

            sigs = [rsi < 35, macd.iloc[-1] > macd_s.iloc[-1], p > sma20, p > sma50, p > sma100, p > sma200, b_pct < 0.15, v_rat > 1.3, smc == "💎 SMC BUY", mc_p > 3, pct_from_ath < -15]
            score = int(sum(1 for s in sigs if s))

            results.append({
                "Nazwa": full_name, "Ticker": t, "Cena": round(p,2), "ATH": round(ath,2), "ATL": round(atl,2), "% od ATH": pct_from_ath,
                f"Dyw. {y1}": div_y1_str, f"Dyw. {y2}": div_y2_str, f"Dyw. {y3}": div_y3_str, "Stopa Dyw. (%)": div_yield,
                "C/Z (P/E)": pe, "Int. C/Z": int_pe, "Forward C/Z": fwd_pe, "C/WK (P/B)": pb,
                "ROE (%)": roe, "Int. ROE": int_roe, "Marża Operac. (%)": op_margin, "Dług/Kapitał": debt_eq, "Wzrost EPS (%)": eps_growth, "Int. EPS": int_eps,
                "RSI": round(rsi,1), "Int. RSI": "🟢 KUPUJ" if rsi < 35 else ("🔴 SPRZEDAJ" if rsi > 65 else "⚪ HOLD"),
                "MACD": round(macd.iloc[-1],2), "Int. MACD": "🟢 KUPUJ" if macd.iloc[-1] > macd_s.iloc[-1] else "🔴 SPRZEDAJ",
                "SMA 20": round(sma20,2), "SMA 50": round(sma50,2), "SMA 100": round(sma100,2), "SMA 200": round(sma200,2),
                "Int. Trend (SMA)": "🟢 SILNA HOSSA" if p > sma200 and p > sma50 else ("🔴 BESSA" if p < sma200 else "⚪ KONSOLIDACJA"),
                "%B (Bollinger)": round(b_pct,2), "Int. Wstęgi": "🟢 ODBICIE (KUP)" if b_pct < 0.15 else "⚪ NEUTRALNIE",
                "Wolumen (Ratio)": round(v_rat,2), "Int. Wolumen": "🟢 AKUMULACJA" if v_rat > 1.3 and delta.iloc[-1] > 0 else "⚪ ZWYKŁY",
                "SMC": smc, "MC Target (30d)": mc_t, "MC Prognoza %": mc_p, "Int. MC": "🟢 WZROST" if mc_p > 3 else "🔴 SPADEK/BOCZNIAK",
                "SUMA BUY": score
            })
        except Exception as e: continue
    
    if not results: return pd.DataFrame()
    return pd.DataFrame(results).sort_values(by='SUMA BUY', ascending=False)

def generate_excel_in_memory(data_dict):
    output = io.BytesIO()
    writer = pd.ExcelWriter(output, engine='xlsxwriter')
    workbook = writer.book
    fmt_green = workbook.add_format({'bg_color': '#C6EFCE', 'font_color': '#006100', 'bold': True})
    fmt_red = workbook.add_format({'bg_color': '#FFC7CE', 'font_color': '#9C0006'})
    fmt_yellow = workbook.add_format({'bg_color': '#FFEB9C', 'font_color': '#9C6500'})
    fmt_header = workbook.add_format({'bg_color': '#2C3E50', 'font_color': '#FFFFFF', 'bold': True, 'border': 1})

    for sheet_name, df in data_dict.items():
        if df.empty: continue
        df.to_excel(writer, sheet_name=sheet_name[:31], index=False)
        ws = writer.sheets[sheet_name[:31]]
        ws.freeze_panes(1, 1)
        ws.autofilter(0, 0, len(df), len(df.columns) - 1)

        for col_num, value in enumerate(df.columns.values):
            ws.write(0, col_num, value, fmt_header)
            col_len = max(df[value].astype(str).map(len).max(), len(value)) + 2
            ws.set_column(col_num, col_num, min(col_len, 25))

        for row in range(1, len(df) + 1):
            for col in range(len(df.columns)):
                cell_val = str(df.iloc[row-1, col])
                if any(x in cell_val for x in ["🟢", "KUP", "BUY", "HOSSA", "WZROST", "AKUMULACJA", "ODBICIE", "TANIO", "WYSOKIE"]):
                    ws.write(row, col, df.iloc[row-1, col], fmt_green)
                elif any(x in cell_val for x in ["🔴", "SPRZEDAJ", "SELL", "BESSA", "SPADEK", "DROGO", "NISKIE"]):
                    ws.write(row, col, df.iloc[row-1, col], fmt_red)
                elif any(x in cell_val for x in ["⚪", "HOLD", "KONSOLIDACJA", "NEUTRALNIE", "ZWYKŁY"]):
                    ws.write(row, col, df.iloc[row-1, col], fmt_yellow)
    writer.close()
    return output.getvalue()


# --- 6. LEWY PANEL NAWIGACYJNY (SIDEBAR) ---
st.sidebar.title("📌 Menu PickPivot")
st.sidebar.markdown(f"Zalogowany jako: **{username if 'username' in locals() else 'DORADCA'}**")

aktywna_zakladka = st.sidebar.radio(
    "Wybierz moduł platformy:",
    [
        "1. Radar Orzecznictwa",
        "2. Ściągacz Interpretacji",
        "3. Global Market Scanner",
        "4. Analiza Wskaźnikowa (XML)",  # ZMIENIONA NAZWA!
        "5. Historia Pobierania (W przyszłości)",
        "6. Ustawienia Systemu (W przyszłości)"
    ]
)

st.sidebar.markdown("---")
if st.sidebar.button("🚪 Wyloguj się", use_container_width=True):
    st.session_state['authenticated'] = False
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.caption("© 2026 PickPivot v13.1 Multi-Tool")

# --- 7. LOGIKA MODUŁÓW ---

if aktywna_zakladka.startswith("1."):
    st.title("⚡ PickPivot: Radar Orzecznictwa")
    st.markdown("Wyszukuje interpretacje podatkowe na podstawie zdefiniowanych słów kluczowych oraz synonimów.")

    konfiguracja = wczytaj_historie(PLIK_KONFIGURACJI_M1)
    przetworzone_id = set(konfiguracja.get("przetworzone_id", []))
    ukonczone_kombinacje = set(konfiguracja.get("ukonczone_kombinacje", []))
    pelne_tresci_cache = wczytaj_pelne_tresci(PLIK_REKORDOW_M1)

    if pelne_tresci_cache:
        st.success(f"💾 BAZA DANYCH RADARU: Zabezpieczono {len(pelne_tresci_cache)} orzeczeń o pełnej treści.")
        colA, colB = st.columns(2)
        with colA:
            if st.button("📄 GENERUJ RAPORT WORD (.docx)", use_container_width=True, type="primary"):
                with st.spinner("Kompilowanie pliku..."):
                    doc = Document()
                    doc.add_heading('Radar PickPivot', 0)
                    for rekord in pelne_tresci_cache:
                        doc.add_heading(f"Sygnatura: {rekord['Sygnatura']}", level=1)
                        doc.add_paragraph(f"Data: {rekord['Data']} | Podatek: {rekord['Podatek']}")
                        doc.add_paragraph(f"Fraza wywołująca: {rekord['Słowo kluczowe']}")
                        doc.add_paragraph(f"Link: {rekord['Link']}")
                        doc.add_heading("Treść:", level=2)
                        doc.add_paragraph(wyczysc_tekst_dla_worda(rekord['Tekst']))
                        doc.add_page_break()
                    output = io.BytesIO()
                    doc.save(output)
                    st.download_button("📥 Pobierz plik", data=output.getvalue(), file_name=f"Radar_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        with colB:
            if st.button("🗑️ Resetuj bazę Radaru", use_container_width=True):
                wyczysc_dane_serwera(PLIK_KONFIGURACJI_M1, PLIK_REKORDOW_M1)
                st.rerun()
        st.markdown("---")

    col1, col2, col3 = st.columns(3)
    with col1: wybrane_lata = st.multiselect("Wybierz lata:", [2024, 2025, 2026])
    with col2: wybrane_miesiace_ui = st.multiselect("Wybierz miesiące:", MIESIACE_PL)
    with col3: wybrane_podatki_ui = st.multiselect("Rodzaj podatku:", ["PIT", "CIT", "VAT", "AKCYZA"], default=["VAT"])

    if st.button("🚀 Uruchom skanowanie słów kluczowych", use_container_width=True):
        if not wybrane_lata or not wybrane_miesiace_ui or not wybrane_podatki_ui:
            st.error("Proszę wybrać komplet parametrów.")
            st.stop()
            
        wybrane_miesiace = [MIESIACE_PL.index(m) + 1 for m in wybrane_miesiace_ui]
            
        pasek_postepu = st.progress(0)
        status_tekst = st.empty()
        calkowita_liczba_zapytan = len(wybrane_lata) * len(wybrane_miesiace) * len(FRAZY_KLUCZOWE) * len(wybrane_podatki_ui)
        zapytania_wykonane = 0
        licznik_trafien = 0

        with requests.Session() as sesja_bazy:
            for rok in wybrane_lata:
                for miesiac in wybrane_miesiace:
                    _, ost_dzien = calendar.monthrange(rok, miesiac)
                    data_start_str = f"{rok}-{miesiac:02d}-01"
                    data_koniec_str = f"{rok}-{miesiac:02d}-{ost_dzien:02d}"
                    for fraza in FRAZY_KLUCZOWE:
                        for podatek in wybrane_podatki_ui:
                            klucz_kombinacji = f"M1_{rok}_{miesiac}_{fraza}_{podatek}"
                            if klucz_kombinacji in ukonczone_kombinacje:
                                zapytania_wykonane += 1
                                continue
                            
                            status_tekst.info(f"Radar odpytuje: {fraza} ({podatek}) dla okresu {miesiac:02d}/{rok}...")
                            lista_trafien, _ = szukaj_w_api_mf(data_start_str, data_koniec_str, fraza, sesja_bazy, podatek, KODY_PODATKOW[podatek])
                            
                            if lista_trafien:
                                aktualne_tresci = wczytaj_pelne_tresci(PLIK_REKORDOW_M1)
                                for dok in lista_trafien:
                                    if dok["id"] not in przetworzone_id:
                                        tekst, status_pobr = pobierz_tekst_pdf(dok["id"])
                                        if tekst:
                                            aktualne_tresci.append({
                                                "Data": dok["data"], "Podatek": dok["typ"], "Sygnatura": dok["sygnatura"],
                                                "Słowo kluczowe": fraza.upper(), "Link": PODGLAD_URL.format(id=dok["id"]), "Tekst": tekst
                                            })
                                            przetworzone_id.add(dok["id"])
                                            konfiguracja["przetworzone_id"].append(dok["id"])
                                            licznik_trafien += 1
                                zapisz_pelne_tresci(PLIK_REKORDOW_M1, aktualne_tresci)
                            
                            ukonczone_kombinacje.add(klucz_kombinacji)
                            konfiguracja["ukonczone_kombinacje"].append(klucz_kombinacji)
                            zapisz_historie(PLIK_KONFIGURACJI_M1, konfiguracja)
                            zapytania_wykonane += 1
                            pasek_postepu.progress(min(1.0, zapytania_wykonane / calkowita_liczba_zapytan))
                            
        status_tekst.success(f"🎉 Zakończono! Zebrano {licznik_trafien} dokumentów.")
        st.balloons()
        time.sleep(3)
        st.rerun()

elif aktywna_zakladka.startswith("2."):
    st.title("📦 Ściągacz Interpretacji")
    st.markdown("Pobiera **wszystkie** interpretacje indywidualne z wybranego okresu.")

    if st.session_state.get('lockout_active_m2', False):
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.error("🔒 SYSTEM OCHRONY PRZED BANEM AKTYWNY")
        elapsed = time.time() - st.session_state.get('lockout_start_m2', 0)
        if elapsed < 300:
            countdown_placeholder = st.empty()
            while elapsed < 300:
                remaining = int(300 - elapsed)
                mins, secs = divmod(remaining, 60)
                countdown_placeholder.markdown(f"<h2 style='text-align: center; color: #ff4b4b;'>⏳ Czas do odblokowania: {mins:02d}:{secs:02d}</h2>", unsafe_allow_html=True)
                time.sleep(1)
                elapsed = time.time() - st.session_state.get('lockout_start_m2', 0)
            st.rerun()
        else:
            st.success("🎉 Gotowe do wznowienia!")
            if st.button("▶️ WZNÓW PRZERWANE POBIERANIE", use_container_width=True, type="primary"):
                st.session_state['lockout_active_m2'] = False
                st.session_state['auto_resume_m2'] = True
                st.rerun()
            st.stop()

    konfiguracja_m2 = wczytaj_historie(PLIK_KONFIGURACJI_M2)
    przetworzone_id_m2 = set(konfiguracja_m2.get("przetworzone_id", []))
    uszkodzone_id_m2 = set(konfiguracja_m2.get("uszkodzone_id", []))
    pelne_tresci_m2 = wczytaj_pelne_tresci(PLIK_REKORDOW_M2)

    if pelne_tresci_m2 or uszkodzone_id_m2:
        st.success(f"💾 BAZA ŚCIĄGACZA: Zabezpieczono {len(pelne_tresci_m2)} dokumentów. (Puste: {len(uszkodzone_id_m2)})")
        if pelne_tresci_m2:
            if st.button("📄 GENERUJ ARCHIWUM WORD (.docx)", use_container_width=True, type="primary"):
                with st.spinner("Składanie dokumentu..."):
                    doc = Document()
                    doc.add_heading('Kompleksowe Archiwum Orzecznictwa', 0)
                    for rekord in pelne_tresci_m2:
                        doc.add_heading(f"Sygnatura: {rekord['Sygnatura']}", level=1)
                        doc.add_paragraph(f"Data: {rekord['Data']} | Podatek: {rekord['Podatek']}")
                        doc.add_paragraph(f"Link: {rekord['Link']}")
                        doc.add_heading("Treść:", level=2)
                        doc.add_paragraph(wyczysc_tekst_dla_worda(rekord['Tekst']))
                        doc.add_page_break()
                    output = io.BytesIO()
                    doc.save(output)
                    st.download_button("📥 Pobierz Archiwum", data=output.getvalue(), file_name=f"Archiwum_Zrzut_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        st.markdown("---")

    col1, col2, col3 = st.columns(3)
    with col1: wybrane_lata_m2 = st.multiselect("Lata:", [2024, 2025, 2026], key="latam2")
    with col2: wybrane_miesiace_m2_ui = st.multiselect("Miesiące:", MIESIACE_PL, key="miesm2")
    with col3: wybrane_podatki_ui_m2 = st.multiselect("Podatki:", ["PIT", "CIT", "VAT", "AKCYZA"], default=["VAT"], key="podm2")

    col_btn1, col_btn2 = st.columns(2)
    with col_btn1: btn_wznow = st.button("▶️ Wznów pobieranie", use_container_width=True)
    with col_btn2: btn_od_nowa = st.button("🔄 Pobierz od nowa", use_container_width=True)

    run_loop = btn_wznow or btn_od_nowa or st.session_state.get('auto_resume_m2', False)

    if run_loop:
        if not st.session_state.get('auto_resume_m2', False):
            if not wybrane_lata_m2 or not wybrane_miesiace_m2_ui or not wybrane_podatki_ui_m2:
                st.error("Wybierz parametry.")
                st.stop()

        if btn_od_nowa:
            wyczysc_dane_serwera(PLIK_KONFIGURACJI_M2, PLIK_REKORDOW_M2)
            konfiguracja_m2 = {"przetworzone_id": [], "ukonczone_kombinacje": [], "uszkodzone_id": []}
            przetworzone_id_m2, uszkodzone_id_m2, pelne_tresci_m2 = set(), set(), []
            if 'queue_m2' in st.session_state: del st.session_state['queue_m2']
            st.toast("🧹 Pamięć wyczyszczona.")

        if btn_wznow and 'queue_m2' in st.session_state:
            del st.session_state['queue_m2']

        status_tekst = st.empty()
        log_szczegolowy = st.empty()

        if st.session_state.get('auto_resume_m2', False) and 'queue_m2' in st.session_state:
            do_pobrania_teraz = st.session_state['queue_m2']
            laczna_liczba_orzeczen = st.session_state.get('laczna_orzeczen_m2', len(do_pobrania_teraz))
            liczba_brakujacych = len(do_pobrania_teraz)
            status_tekst.success("▶️ Kontynuuję...")
        else:
            status_tekst.info("🔍 Krok 1/2: Zliczanie...")
            wszystkie_orzeczenia_w_mf, do_pobrania_teraz = [], []
            wybrane_miesiace_m2 = [MIESIACE_PL.index(m) + 1 for m in wybrane_miesiace_m2_ui]

            with requests.Session() as sesja_bazy:
                for rok in wybrane_lata_m2:
                    for miesiac in wybrane_miesiace_m2:
                        _, ost_dzien = calendar.monthrange(rok, miesiac)
                        data_start_str = f"{rok}-{miesiac:02d}-01"
                        data_koniec_str = f"{rok}-{miesiac:02d}-{ost_dzien:02d}"
                        for podatek in wybrane_podatki_ui_m2:
                            lista_trafien, _ = pobierz_wszystko_z_okresu(data_start_str, data_koniec_str, sesja_bazy, podatek, KODY_PODATKOW[podatek])
                            for dok in lista_trafien:
                                wszystkie_orzeczenia_w_mf.append(dok)
                                if dok["id"] not in przetworzone_id_m2 and dok["id"] not in uszkodzone_id_m2:
                                    do_pobrania_teraz.append(dok)

            laczna_liczba_orzeczen = len(wszystkie_orzeczenia_w_mf)
            liczba_brakujacych = len(do_pobrania_teraz)
            st.session_state['laczna_orzeczen_m2'] = laczna_liczba_orzeczen

        st.session_state['auto_resume_m2'] = False
        st.info(f"Odnaleziono: {laczna_liczba_orzeczen}. Do pobrania: {liczba_brakujacych}.")
        
        if liczba_brakujacych == 0:
            status_tekst.success("✔️ Pula wyczerpana.")
            st.stop()

        time.sleep(1)
        status_tekst.info(f"⏳ Krok 2/2: Pobieranie (0/{liczba_brakujacych})...")
        pasek_postepu = st.progress(0)
        
        licznik_pobranych_w_sesji, licznik_uszkodzonych_w_sesji = 0, 0
        aktualne_tresci_m2 = wczytaj_pelne_tresci(PLIK_REKORDOW_M2)

        for idx, dok in enumerate(do_pobrania_teraz):
            st.session_state['queue_m2'] = do_pobrania_teraz[idx:]
            tekst, status_pobr = pobierz_tekst_pdf(dok["id"])
            
            if tekst:
                aktualne_tresci_m2.append({"Data": dok["data"], "Podatek": dok["typ"], "Sygnatura": dok["sygnatura"], "Link": PODGLAD_URL.format(id=dok["id"]), "Tekst": tekst})
                przetworzone_id_m2.add(dok["id"])
                konfiguracja_m2["przetworzone_id"].append(dok["id"])
                licznik_pobranych_w_sesji += 1
            else:
                if status_pobr in ["BRAK_PLIKU", "BŁĄD_CZYTANIA"]:
                    uszkodzone_id_m2.add(dok["id"])
                    konfiguracja_m2["uszkodzone_id"].append(dok["id"])
                    licznik_uszkodzonych_w_sesji += 1
                else:
                    st.session_state['lockout_active_m2'] = True
                    st.session_state['lockout_start_m2'] = time.time()
                    st.session_state['queue_m2'] = do_pobrania_teraz[idx:]
                    zapisz_pelne_tresci(PLIK_REKORDOW_M2, aktualne_tresci_m2)
                    zapisz_historie(PLIK_KONFIGURACJI_M2, konfiguracja_m2)
                    st.rerun()
                
            zapisz_pelne_tresci(PLIK_REKORDOW_M2, aktualne_tresci_m2)
            zapisz_historie(PLIK_KONFIGURACJI_M2, konfiguracja_m2)

            status_tekst.info(f"⏳ Pobrane: {licznik_pobranych_w_sesji} | Puste: {licznik_uszkodzonych_w_sesji}")
            pasek_postepu.progress((idx + 1) / liczba_brakujacych)
            time.sleep(random.uniform(1.5, 2.5))

        if 'queue_m2' in st.session_state: del st.session_state['queue_m2']
        status_tekst.success(f"🎉 SUKCES! Pobrano {licznik_pobranych_w_sesji}.")
        st.balloons()
        time.sleep(4)
        st.rerun()

elif aktywna_zakladka.startswith("3."):
    st.title("📊 PickPivot: Global Market Scanner")
    st.markdown("Weryfikacja danych technicznych oraz fundamentalnych giełd z Europy i USA.")

    ALL_MARKETS = {"Polska (WIG)": {**WIG20_MAP, **MWIG40_MAP}, "Niemcy (DAX)": DAX_MAP, "Francja (CAC 40)": CAC40_MAP, "UK (FTSE 100)": FTSE_MAP, "Hiszpania (IBEX 35)": IBEX_MAP, "Szwecja (OMX 30)": OMX_MAP, "Norwegia (OBX)": OBX_MAP, "USA (S&P 500)": get_sp500_map()}
    selected_markets = st.sidebar.multiselect("Wybierz rynki do analizy:", options=list(ALL_MARKETS.keys()), default=["Polska (WIG)"])

    if st.button("🚀 Uruchom Analizę Rynków", type="primary", use_container_width=True):
        if not selected_markets:
            st.warning("Proszę wybrać przynajmniej jeden rynek z panelu bocznego.")
        else:
            st.info("Pobieranie danych z Yahoo Finance...")
            final_results = {}
            progress_bar = st.progress(0)
            
            for i, market_name in enumerate(selected_markets):
                final_results[market_name] = analyze_market(ALL_MARKETS[market_name], market_name)
                progress_bar.progress((i + 1) / len(selected_markets))
                
            st.success("🎉 Analiza zakończona pomyślnie!")
            
            st.subheader("Wyniki Skanowania")
            tabs = st.tabs(list(final_results.keys()))
            for tab, market_name in zip(tabs, final_results.keys()):
                with tab:
                    df = final_results[market_name]
                    if not df.empty: st.dataframe(df, use_container_width=True, hide_index=True)
                    else: st.write("Brak danych.")
                        
            st.markdown("---")
            st.subheader("Eksport i Dystrybucja")
            excel_data = generate_excel_in_memory(final_results)
            
            colA, colB = st.columns(2)
            with colA:
                st.download_button("📥 Pobierz plik Excel", data=excel_data, file_name=f"Global_Report_{datetime.now().strftime('%Y%m%d')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            with colB:
                DISCORD_WEBHOOK_URL = "TUTAJ_WKLEJ_SWÓJ_LINK_WEBHOOK_Z_DISCORDA"
                if st.button("💬 Wyślij na Discord", use_container_width=True):
                    if DISCORD_WEBHOOK_URL == "TUTAJ_WKLEJ_SWÓJ_LINK_WEBHOOK_Z_DISCORDA":
                        st.error("Wklej link Webhook w kodzie!")
                    else:
                        with st.spinner("Wysyłanie..."):
                            payload = {"content": f"📊 **Nowy Raport PickPivot Scanner** ({datetime.now().strftime('%Y-%m-%d %H:%M')}) gotowy!"}
                            files = {"file": (f"Global_Report_{datetime.now().strftime('%Y%m%d')}.xlsx", excel_data, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")}
                            try:
                                resp = requests.post(DISCORD_WEBHOOK_URL, data=payload, files=files)
                                if resp.status_code in [200, 204]:
                                    st.success("✔️ Wysłano na Discord!")
                                    st.balloons()
                                else: st.error("❌ Błąd wysyłania.")
                            except: st.error("❌ Błąd sieciowy.")

# --- NOWOŚĆ: WYWOŁANIE MODUŁU 4 Z OSOBNEGO PLIKU ---
elif aktywna_zakladka.startswith("4."):
    import cfo_analyzer
    cfo_analyzer.run_module()

else:
    st.title(f"🛠️ {aktywna_zakladka}")
    st.info("Ta funkcjonalność jest obecnie w fazie projektowania i zostanie dodana w przyszłości.", icon="ℹ️")
