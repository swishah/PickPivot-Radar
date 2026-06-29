"""
downloader.py - Sciagacz Interpretacji v3
Backend archiwum wykrywany LAZILY przy pierwszym wywolaniu run_module().
"""

import streamlit as st
import time
import io
import requests
import calendar
from docx import Document
from datetime import datetime
import utils

CZAS_LOCKOUT_S = 300


# ---------------------------------------------------------------------------
# LAZY LOADING BACKENDU — wywolywane wewnatrz run_module(), nie przy imporcie
# ---------------------------------------------------------------------------
def _get_archiwum():
    """
    Zwraca (modul_archiwum, nazwa) lub (None, None).
    Cachowane w session_state zeby nie importowac przy kazdym rerunie.
    """
    if "arch_mod" not in st.session_state:
        mod, nazwa = _wykryj_archiwum()
        st.session_state["arch_mod"]   = mod
        st.session_state["arch_nazwa"] = nazwa
    return st.session_state["arch_mod"], st.session_state["arch_nazwa"]


def _wykryj_archiwum():
    try:
        sec = st.secrets
        if "supabase" in sec:
            s = sec["supabase"]
            ma_host = bool(s.get("host","")) and bool(s.get("user","")) and bool(s.get("password",""))
            ma_url  = str(s.get("url","")).startswith("postgresql")
            if ma_host or ma_url:
                import archiwum_supabase as _arch
                return _arch, "Supabase (PostgreSQL)"
    except Exception:
        pass
    try:
        sec = st.secrets
        if "google_drive" in sec:
            if sec["google_drive"].get("folder_id",""):
                import archiwum_drive as _arch
                return _arch, "Google Drive"
    except Exception:
        pass
    return None, None


# ---------------------------------------------------------------------------
# GENEROWANIE WORD
# ---------------------------------------------------------------------------
def _generuj_word(rekordy: list, filtry_opis: str = "") -> bytes:
    doc = Document()
    doc.add_heading("PickPivot - Archiwum Interpretacji Podatkowych", 0)
    doc.add_paragraph(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    if filtry_opis:
        doc.add_paragraph(f"Zakres: {filtry_opis}")
    doc.add_paragraph(f"Liczba dokumentow: {len(rekordy)}")
    doc.add_page_break()
    for r in sorted(rekordy, key=lambda x: x.get("Data",""), reverse=True):
        doc.add_heading(f"Sygnatura: {r['Sygnatura']}", 1)
        doc.add_paragraph(f"Data:    {r['Data']}")
        doc.add_paragraph(f"Podatek: {r['Podatek']}")
        doc.add_paragraph(f"Link:    {r['Link']}")
        if r.get("Format"):
            doc.add_paragraph(f"Zrodlo:  {r['Format']}")
        doc.add_paragraph(utils.wyczysc_tekst_dla_worda(r["Tekst"]))
        doc.add_page_break()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PANEL ARCHIWUM
# ---------------------------------------------------------------------------
def _renderuj_panel_archiwum(arch, nazwa):
    st.markdown("### Wspoldzielone Archiwum Interpretacji")

    if arch is None:
        st.warning(
            "Archiwum nie jest skonfigurowane. "
            "Dodaj sekcje [supabase] lub [google_drive] w Streamlit Secrets. "
            "Pobieranie nadal dziala, ale dokumenty nie beda zapamietywane miedzy sesjami."
        )
        st.markdown("---")
        return

    st.caption(
        f"Backend: **{nazwa}** | "
        "Dokumenty pobrane przez kogokolwiek sa dostepne dla wszystkich uzytkownikow."
    )

    with st.spinner(f"Laczenie z {nazwa}..."):
        try:
            stats = arch.statystyki_archiwum()
        except Exception as e:
            st.error(f"Blad polaczenia z archiwum: {e}")
            st.markdown("---")
            return

    if not stats.get("polaczenie"):
        st.error(f"Brak polaczenia z {nazwa}. Sprawdz konfiguracje w Streamlit Secrets.")
        st.markdown("---")
        return

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Dokumentow w archiwum", f"{stats['total']:,}".replace(",", " "))
    c2.metric("Ukonczonych kombinacji", stats["ukonczone_kombinacje"])
    c3.metric("Ostatnie pobranie",      stats["ostatnie_pobranie"])
    c4.metric("Backend",                nazwa)

    if stats.get("per_podatek"):
        st.caption("  |  ".join(f"**{k}:** {v}" for k, v in stats["per_podatek"].items()))

    with st.expander("Przegladaj i pobierz z archiwum"):
        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            filtr_pod = st.multiselect("Podatek:", ["PIT","CIT","VAT","AKCYZA"], key="arch_pod")
        with col_f2:
            filtr_rok = st.selectbox("Rok:", [None,2024,2025,2026],
                format_func=lambda x: "Wszystkie" if x is None else str(x), key="arch_rok")
        with col_f3:
            filtr_mies = st.selectbox("Miesiac:", [None]+list(range(1,13)),
                format_func=lambda x: "Wszystkie" if x is None else utils.MIESIACE_PL[x-1],
                key="arch_mies")

        if st.button("Szukaj w archiwum", use_container_width=True):
            with st.spinner("Szukam..."):
                wyniki = []
                for pod in (filtr_pod or [None]):
                    wyniki += arch.pobierz_rekordy_z_archiwum(
                        podatek=pod, rok=filtr_rok, miesiac=filtr_mies)
                seen, unikalne = set(), []
                for r in wyniki:
                    rid = arch._id_z_rekordu(r)
                    if rid not in seen:
                        seen.add(rid); unikalne.append(r)
            st.session_state["arch_podglad"] = unikalne

        podglad = st.session_state.get("arch_podglad", [])
        if podglad:
            st.success(f"Znaleziono **{len(podglad)}** dokumentow.")
            for r in podglad[:20]:
                st.markdown(f"- **{r['Sygnatura']}** | {r['Data']} | {r['Podatek']}")
            if len(podglad) > 20:
                st.caption(f"... i {len(podglad)-20} wiecej.")
            opis = ", ".join(filter(None, [
                "/".join(filtr_pod) if filtr_pod else None,
                str(filtr_rok) if filtr_rok else None,
                utils.MIESIACE_PL[filtr_mies-1] if filtr_mies else None,
            ])) or "cale archiwum"
            st.download_button("Pobierz jako Word", data=_generuj_word(podglad, opis),
                file_name=f"Archiwum_{opis.replace(', ','_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True)
    st.markdown("---")


# ---------------------------------------------------------------------------
# EKSPORT DO WORD
# ---------------------------------------------------------------------------
def _pobierz_i_wyeksportuj(arch, wybrane_lata, wybrane_miesiace_ui, wybrane_podatki):
    mies_num = [utils.MIESIACE_PL.index(m)+1 for m in (wybrane_miesiace_ui or [])]
    rekordy  = []
    with st.spinner("Pobieram z archiwum..."):
        for pod in (wybrane_podatki or [None]):
            for rok in (wybrane_lata or [None]):
                for mies in (mies_num or [None]):
                    rekordy += arch.pobierz_rekordy_z_archiwum(podatek=pod, rok=rok, miesiac=mies)
    seen, unikalne = set(), []
    for r in rekordy:
        rid = arch._id_z_rekordu(r)
        if rid not in seen:
            seen.add(rid); unikalne.append(r)
    if not unikalne:
        st.warning("Brak dokumentow w archiwum dla wybranych filtrow."); return
    opis = ", ".join(filter(None, [
        "/".join(wybrane_podatki) if wybrane_podatki else None,
        "/".join(str(r) for r in wybrane_lata) if wybrane_lata else None,
        "/".join(wybrane_miesiace_ui) if wybrane_miesiace_ui else None,
    ])) or "pelne archiwum"
    st.download_button(
        f"Pobierz Word ({len(unikalne)} interpretacji)",
        data=_generuj_word(unikalne, opis),
        file_name=f"Interpretacje_{opis.replace(', ','_')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        type="primary", use_container_width=True)


# ---------------------------------------------------------------------------
# GLOWNY MODUL
# ---------------------------------------------------------------------------
def run_module():
    st.title("Sciagacz Interpretacji v3")

    # Blokada anty-ban
    if st.session_state.get("lockout_active_m2", False):
        uplynelo  = time.time() - st.session_state.get("lockout_start_m2", 0)
        pozostalo = max(0, int(CZAS_LOCKOUT_S - uplynelo))
        st.error(f"OCHRONA ANTY-BAN — Przerwa ochronna ({pozostalo}s pozostalo)")
        st.progress(min(1.0, uplynelo / CZAS_LOCKOUT_S))
        if uplynelo < CZAS_LOCKOUT_S:
            time.sleep(1); st.rerun()
        else:
            if st.button("WZNOW", type="primary", use_container_width=True):
                st.session_state["lockout_active_m2"] = False
                st.session_state["auto_resume_m2"]    = True
                st.rerun()
        st.stop()

    # Wykryj backend TUTAJ - po zaladowaniu Streamlit, z dostepem do st.secrets
    arch, nazwa = _get_archiwum()

    _renderuj_panel_archiwum(arch, nazwa)

    st.markdown("### Pobierz nowe interpretacje z MF")
    col1, col2, col3 = st.columns(3)
    with col1: wybrane_lata = st.multiselect("Lata:", [2024,2025,2026], key="l2")
    with col2: wybrane_miesiace_ui = st.multiselect("Miesiace:", utils.MIESIACE_PL, key="m2")
    with col3: wybrane_podatki = st.multiselect("Podatki:", ["PIT","CIT","VAT","AKCYZA"], default=["VAT"], key="p2")

    with st.expander("Ustawienia pobierania"):
        workers = st.slider("Rownoległe watki", 1, 8, 5)
        st.session_state["m2_workers"] = workers

    c1, c2 = st.columns(2)
    with c1: btn_start = st.button("Uruchom pobieranie", type="primary", use_container_width=True)
    with c2: btn_wznow = st.button("Wznow przerywane", use_container_width=True)

    if not (btn_start or btn_wznow or st.session_state.get("auto_resume_m2", False)):
        st.stop()

    konfig        = utils.wczytaj_historie(utils.PLIK_KONFIGURACJI_M2)
    uszkodzone_id = set(konfig.get("uszkodzone_id", []))

    if arch is not None:
        with st.spinner("Sprawdzam archiwum..."):
            id_w_archiwum        = arch.pobierz_id_z_archiwum()
            ukonczone_kombinacje = arch.pobierz_ukonczone_kombinacje()
    else:
        id_w_archiwum, ukonczone_kombinacje = set(), set()

    wszystkie_znane = id_w_archiwum | uszkodzone_id
    bledy_api = 0

    if st.session_state.get("auto_resume_m2", False) and "queue_m2" in st.session_state:
        do_pobrania = st.session_state["queue_m2"]
        st.info(f"Wznawianie — {len(do_pobrania)} dokumentow w kolejce.")
    else:
        if not wybrane_lata or not wybrane_miesiace_ui or not wybrane_podatki:
            st.error("Wybierz przynajmniej jeden rok, miesiac i podatek."); st.stop()

        wybrane_miesiace = [utils.MIESIACE_PL.index(m)+1 for m in wybrane_miesiace_ui]
        wszystkie_kombinacje = [
            (rok, mies, pod)
            for rok  in wybrane_lata
            for mies in wybrane_miesiace
            for pod  in wybrane_podatki
        ]

        if arch is not None:
            do_api = [k for k in wszystkie_kombinacje
                      if arch._klucz_kombinacji(k[2],k[0],k[1]) not in ukonczone_kombinacje]
            z_arch = len(wszystkie_kombinacje) - len(do_api)
            if z_arch:
                st.success(f"{z_arch} kombinacji juz w archiwum — pominiete.")
        else:
            do_api = wszystkie_kombinacje

        do_pobrania, znalezione_ogolem = [], 0

        if do_api:
            kontener_st = st.empty()
            pasek_list  = st.progress(0)
            with requests.Session() as sesja:
                for i, (rok, mies, pod) in enumerate(do_api, 1):
                    _, ost = calendar.monthrange(rok, mies)
                    kontener_st.info(f"[{i}/{len(do_api)}] Odpytuje MF: {utils.MIESIACE_PL[mies-1]} {rok} ({pod})...")
                    lista, status = utils.pobierz_wszystko_z_okresu(
                        f"{rok}-{mies:02d}-01", f"{rok}-{mies:02d}-{ost:02d}",
                        sesja, pod, utils.KODY_PODATKOW[pod])
                    if status in ("ERROR","TIMEOUT"): bledy_api += 1
                    znalezione_ogolem += len(lista)
                    for d in lista:
                        if d["id"] not in wszystkie_znane:
                            do_pobrania.append(d)
                    pasek_list.progress(i / len(do_api))
            kontener_st.empty(); pasek_list.empty()

        if bledy_api:
            st.error(f"Serwer MF odrzucil {bledy_api} zapytan. Odczekaj 15-30 min."); st.stop()

        st.info(f"Znaleziono **{znalezione_ogolem}** w MF. W archiwum: **{len(id_w_archiwum)}**. Do pobrania: **{len(do_pobrania)}**.")

    st.session_state["auto_resume_m2"] = False

    if not do_pobrania:
        st.success("Wszystko w archiwum! Generuje plik Word...")
        if arch is not None:
            _pobierz_i_wyeksportuj(arch, wybrane_lata, wybrane_miesiace_ui, wybrane_podatki)
        st.stop()

    workers_count = st.session_state.get("m2_workers", 5)
    st.markdown(f"### Pobieranie {len(do_pobrania)} dokumentow ({workers_count} watkow)")
    pasek_gl = st.progress(0); k_licznik = st.empty(); k_log = st.empty()
    log_lines = []; stan = {"ok": 0, "blad": 0}

    def on_postep(completed, total, sygnatura, status):
        with utils._lock:
            if   status == "OK":                            stan["ok"]   += 1; ikona = "OK"
            elif status in ("BRAK_PLIKU","BLAD_CZYTANIA"): stan["blad"] += 1; ikona = "BRAK"
            elif status == "BLOKADA":                       ikona = "BLOK"
            elif status == "POMINIETY":                     ikona = "POMIN"
            else:                                           ikona = "?"
            log_lines.insert(0, f"{ikona} [{completed}/{total}] {sygnatura}")
            if len(log_lines) > 8: log_lines.pop()
        pasek_gl.progress(completed / total)
        k_licznik.markdown(f"Pobrano: **{stan['ok']}** | Bledy: **{stan['blad']}** | Lacznie: **{completed}/{total}**")
        k_log.code("\n".join(log_lines), language=None)

    st.session_state["queue_m2"] = do_pobrania
    nowe_tresci, nowe_przetworzone, nowe_uszkodzone, blokada = \
        utils.pobierz_dokumenty_rownolegle(
            do_pobrania, wszystkie_znane, uszkodzone_id,
            callback_postep=on_postep, workers=workers_count)

    if nowe_tresci and arch is not None:
        with st.spinner(f"Zapisuje do archiwum ({nazwa})..."):
            nowych = arch.zapisz_wiele_do_archiwum(nowe_tresci, "downloader_v3")
        st.info(f"Zapisano **{nowych}** nowych rekordow do archiwum.")

    if arch is not None and not blokada and wybrane_lata and wybrane_miesiace_ui and wybrane_podatki:
        mies_num2 = [utils.MIESIACE_PL.index(m)+1 for m in wybrane_miesiace_ui]
        for rok in wybrane_lata:
            for mies in mies_num2:
                for pod in wybrane_podatki:
                    arch.oznacz_kombinacje(pod, rok, mies)

    konfig["przetworzone_id"] = list(set(konfig.get("przetworzone_id",[])) | set(nowe_przetworzone))
    if not blokada:
        konfig["uszkodzone_id"] = list(set(konfig.get("uszkodzone_id",[])) | set(nowe_uszkodzone))
    utils.zapisz_historie(utils.PLIK_KONFIGURACJI_M2, konfig)

    if blokada:
        pobrane_id   = set(nowe_przetworzone)
        kolejka_rest = [d for d in do_pobrania
                        if d["id"] not in pobrane_id and d["id"] not in set(nowe_uszkodzone)]
        st.session_state["queue_m2"]          = kolejka_rest
        st.session_state["lockout_active_m2"] = True
        st.session_state["lockout_start_m2"]  = time.time()
        st.rerun()

    st.session_state.pop("queue_m2", None)
    pasek_gl.progress(1.0)
    st.success(f"Zakończono! Pobrano **{len(nowe_tresci)}** nowych dokumentow.")
    if arch is not None:
        _pobierz_i_wyeksportuj(arch, wybrane_lata, wybrane_miesiace_ui, wybrane_podatki)
    st.balloons()
