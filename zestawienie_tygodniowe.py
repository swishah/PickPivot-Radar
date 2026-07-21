# -*- coding: utf-8 -*-
"""
MODUŁ: Zestawienie Tygodniowe — Skaner Doradca
===============================================================================
Cztery zakładki (PIT / CIT / VAT / AKCYZA).

OBIEG:
  1. Wgrywasz plik z GPT „Tygodniowy Research” (DOCX zalecany; PDF działa
     najlepiej-jak-się-da). NIE wybierasz tygodnia przy wgrywaniu.
  2. Moduł parsuje tabelę pliku (L.p. | Sygnatura | Data wydania | Temat |
     Streszczenie) i KAŻDĄ interpretację przypisuje do tygodnia JEJ daty
     wydania (tydzień kalendarzowy pn–nd).
  3. Wybierasz tydzień — moduł składa tabelę wszystkich interpretacji z tego
     tygodnia.

PUBLIKACJE OPÓŹNIONE (zasada uzgodniona wcześniej):
  * „Okres researchu” = tydzień najświeższej interpretacji w pliku.
  * Interpretacja, której data wydania należy do WCZEŚNIEJSZEGO tygodnia niż
    okres researchu, trafia do swojego właściwego (wcześniejszego) tygodnia,
    a w tygodniu researchu pokazywana jest DODATKOWO — podświetlona na zielono
    („publikacja opóźniona”).

DANE: tabela interpretacje_streszczenia (tworzona automatycznie).
Tydzień jako klucz ISO „RRRR-Www” (porównania łańcuchowe są chronologiczne,
bo rok i numer tygodnia są wyzerowane do stałej szerokości).
===============================================================================
"""

from __future__ import annotations

import datetime as dt
import hashlib
import html
import io
import re

import streamlit as st

import archiwum_supabase
import auth
from streszczacz_openrouter import _waliduj_branze, _waliduj_przedmioty

# ---------------------------------------------------------------------------
# KONFIGURACJA
# ---------------------------------------------------------------------------
PODATKI = ["PIT", "CIT", "VAT", "AKCYZA"]

ZIELEN_GLOWNA = "#386520"
ZIELEN_TLO = "#dcefd8"
ZIELEN_TEKST = "#1b3d0f"


# ---------------------------------------------------------------------------
# POŁĄCZENIE (to samo co reszta aplikacji)
# ---------------------------------------------------------------------------
def _zapytaj(sql: str, parametry: tuple | None = None) -> list[dict]:
    return archiwum_supabase._get_db().wykonaj(sql, parametry, fetch=True)


def _wykonaj(sql: str, parametry: tuple | None = None) -> int:
    return archiwum_supabase._get_db().wykonaj(sql, parametry, fetch=False)


@st.cache_resource(show_spinner=False)
def _zapewnij_tabele() -> bool:
    _wykonaj(
        """
        CREATE TABLE IF NOT EXISTS interpretacje_streszczenia (
            id               SERIAL PRIMARY KEY,
            podatek          TEXT NOT NULL,
            sygnatura        TEXT NOT NULL,
            data_wyd         TEXT NOT NULL,     -- YYYY-MM-DD
            tydzien_wydania  TEXT NOT NULL,     -- RRRR-Www (tydzień daty wydania)
            tydzien_research TEXT NOT NULL,     -- RRRR-Www (okres pliku researchu)
            temat            TEXT DEFAULT '',
            branza           TEXT DEFAULT '',
            przedmiot        TEXT DEFAULT '',
            streszczenie     TEXT DEFAULT '',
            nazwa_pliku      TEXT DEFAULT '',
            wgrano           TEXT NOT NULL,
            UNIQUE (podatek, sygnatura)
        )
        """
    )
    _wykonaj(
        "ALTER TABLE interpretacje_streszczenia "
        "ADD COLUMN IF NOT EXISTS branza TEXT DEFAULT ''"
    )
    _wykonaj(
        "ALTER TABLE interpretacje_streszczenia "
        "ADD COLUMN IF NOT EXISTS przedmiot TEXT DEFAULT ''"
    )
    _wykonaj(
        "CREATE INDEX IF NOT EXISTS idx_is_tyg "
        "ON interpretacje_streszczenia (podatek, tydzien_wydania)"
    )
    return True


# ---------------------------------------------------------------------------
# TYGODNIE — pn–nd (pełny tydzień kalendarzowy = tydzień ISO)
# ---------------------------------------------------------------------------
def _poniedzialek(d: dt.date) -> dt.date:
    return d - dt.timedelta(days=d.weekday())


def _niedziela(d: dt.date) -> dt.date:
    return _poniedzialek(d) + dt.timedelta(days=6)


def _klucz_tygodnia(d: dt.date) -> str:
    rok, nr, _ = d.isocalendar()
    return f"{rok}-W{nr:02d}"


def _monday_z_klucza(klucz: str) -> dt.date | None:
    try:
        rok, wk = klucz.split("-W")
        return dt.date.fromisocalendar(int(rok), int(wk), 1)
    except Exception:
        return None


def _etykieta_tygodnia(klucz: str) -> str:
    pon = _monday_z_klucza(klucz)
    if not pon:
        return klucz
    nd = pon + dt.timedelta(days=6)
    return f"Tydzień {pon.isocalendar()[1]:02d}/{pon.year}  ·  {pon:%d.%m} – {nd:%d.%m.%Y} (pn–nd)"


# ---------------------------------------------------------------------------
# PARSOWANIE PLIKU
# ---------------------------------------------------------------------------
def _parsuj_date(s: str) -> dt.date | None:
    s = (s or "").strip()
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y"):
        try:
            return dt.datetime.strptime(s[:10], fmt).date()
        except Exception:
            pass
    m = re.search(r"(\d{2})[.\-/](\d{2})[.\-/](\d{4})", s)
    if m:
        try:
            return dt.date(int(m[3]), int(m[2]), int(m[1]))
        except Exception:
            return None
    m = re.search(r"(\d{4})-(\d{2})-(\d{2})", s)
    if m:
        try:
            return dt.date(int(m[1]), int(m[2]), int(m[3]))
        except Exception:
            return None
    return None


def _parsuj_docx(bajty: bytes) -> list[dict]:
    from docx import Document
    doc = Document(io.BytesIO(bajty))
    if not doc.tables:
        return []
    tab = doc.tables[0]
    naglowki = [c.text.strip().lower() for c in tab.rows[0].cells]

    def idx(*klucze):
        for i, h in enumerate(naglowki):
            if any(k in h for k in klucze):
                return i
        return None

    i_syg = idx("sygnatura", "znak")
    i_dat = idx("data")
    i_tem = idx("temat", "przedmiot")
    i_str = idx("streszczenie", "omówienie", "omowienie", "opis")
    i_lp = idx("l.p", "lp", "l. p")
    i_br = idx("branż", "branza")
    i_prz = idx("przedmiot", "obszar")

    def g(c, i):
        return c[i].strip() if (i is not None and i < len(c)) else ""

    wiersze = []
    for row in tab.rows[1:]:
        c = [x.text for x in row.cells]
        syg, dat = g(c, i_syg), g(c, i_dat)
        if not syg and not dat:
            continue
        wiersze.append({
            "lp": g(c, i_lp),
            "sygnatura": syg,
            "data_raw": dat,
            "temat": g(c, i_tem),
            "streszczenie": g(c, i_str),
            "branza": g(c, i_br),
            "przedmiot": g(c, i_prz),
        })
    return wiersze


def _parsuj_pdf(bajty: bytes) -> list[dict]:
    """Best-effort dla PDF: wyłapuje wiersze zaczynające się sygnaturą i datą.
    DOCX jest znacznie pewniejszy — PDF traktujemy jako awaryjny."""
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(bajty))
    tekst = "\n".join((s.extract_text() or "") for s in reader.pages)
    wiersze = []
    wzor = re.compile(
        r"([0-9]{3,4}-[A-ZĄĆĘŁŃÓŚŹŻ0-9.\-]+)\s+"       # sygnatura
        r"(\d{2}[.\-/]\d{2}[.\-/]\d{4})\s+"            # data
        r"(.+?)(?=(?:[0-9]{3,4}-[A-ZĄĆĘŁŃÓŚŹŻ0-9.\-]+\s+\d{2}[.\-/]\d{2}[.\-/]\d{4})|$)",
        re.S,
    )
    for m in wzor.finditer(tekst):
        reszta = re.sub(r"\s+", " ", m[3]).strip()
        wiersze.append({
            "lp": "", "sygnatura": m[1].strip(), "data_raw": m[2].strip(),
            "temat": "", "streszczenie": reszta,
        })
    return wiersze


def _parsuj_plik(nazwa: str, bajty: bytes) -> list[dict]:
    if nazwa.lower().endswith(".docx"):
        return _parsuj_docx(bajty)
    if nazwa.lower().endswith(".pdf"):
        return _parsuj_pdf(bajty)
    return []


# ---------------------------------------------------------------------------
# ZAPIS PO PARSOWANIU — przypisanie do tygodni
# ---------------------------------------------------------------------------
def _zapisz_interpretacje(podatek: str, nazwa: str, wiersze: list[dict]) -> tuple[int, int, str | None]:
    """Zwraca (zapisane, pominiete_bez_daty, tydzien_research_klucz)."""
    daty = [_parsuj_date(w["data_raw"]) for w in wiersze]
    prawidlowe = [d for d in daty if d]
    if not prawidlowe:
        return 0, len(wiersze), None

    tydzien_research = _klucz_tygodnia(max(prawidlowe))  # tydzień najświeższej
    teraz = dt.datetime.now().isoformat(timespec="seconds")

    zapisane, pominiete = 0, 0
    for w, d in zip(wiersze, daty):
        if not d:
            pominiete += 1
            continue
        _wykonaj(
            """
            INSERT INTO interpretacje_streszczenia
                (podatek, sygnatura, data_wyd, tydzien_wydania, tydzien_research,
                 temat, streszczenie, branza, przedmiot, nazwa_pliku, wgrano)
            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            ON CONFLICT (podatek, sygnatura) DO UPDATE SET
                data_wyd         = EXCLUDED.data_wyd,
                tydzien_wydania  = EXCLUDED.tydzien_wydania,
                tydzien_research = EXCLUDED.tydzien_research,
                temat            = EXCLUDED.temat,
                streszczenie     = EXCLUDED.streszczenie,
                branza           = EXCLUDED.branza,
                przedmiot        = EXCLUDED.przedmiot,
                nazwa_pliku      = EXCLUDED.nazwa_pliku,
                wgrano           = EXCLUDED.wgrano
            """,
            (podatek, w["sygnatura"], d.isoformat(), _klucz_tygodnia(d),
             tydzien_research, w.get("temat", ""), w.get("streszczenie", ""),
             ", ".join(_waliduj_branze(w.get("branza", ""))),
             "; ".join(_waliduj_przedmioty(w.get("przedmiot", ""), podatek)),
             nazwa, teraz),
        )
        zapisane += 1
    return zapisane, pominiete, tydzien_research


# ---------------------------------------------------------------------------
# ODCZYT — tygodnie z danymi + interpretacje danego tygodnia
# ---------------------------------------------------------------------------
def _tygodnie_z_danymi(podatek: str) -> list[str]:
    rows = _zapytaj(
        "SELECT tydzien_wydania AS w FROM interpretacje_streszczenia WHERE podatek=%s "
        "UNION SELECT tydzien_research FROM interpretacje_streszczenia WHERE podatek=%s",
        (podatek, podatek),
    )
    return sorted({r["w"] for r in rows if r.get("w")}, reverse=True)


LIMIT_WIERSZY_5 = 50
SORT_KOLUMNY_5 = {"Data wydania": "data_wyd", "Sygnatura": "sygnatura"}


def _interpretacje_sortowane(podatek: str, sort_kol: str,
                             malejaco: bool) -> list[dict]:
    """50 wgranych interpretacji danego podatku, z wybranym sortowaniem.
    Bez daty publikacji — w module 5 byłaby zawsze datą wgrania pliku."""
    kol = SORT_KOLUMNY_5.get(sort_kol, "data_wyd")
    kier = "DESC" if malejaco else "ASC"
    return _zapytaj(
        f"""
        SELECT sygnatura, data_wyd, temat, streszczenie,
               COALESCE(branza, '') AS branza,
               COALESCE(przedmiot, '') AS przedmiot
        FROM interpretacje_streszczenia
        WHERE podatek = %s
        ORDER BY {kol} {kier} NULLS LAST, sygnatura
        LIMIT {LIMIT_WIERSZY_5}
        """,
        (podatek,),
    )


# ---------------------------------------------------------------------------
# RENDER TABELI (HTML — pełna kontrola nad zawijaniem i zielonym oznaczeniem)
# ---------------------------------------------------------------------------
def _kolory():
    try:
        import paleta
        p = paleta.paleta()
        return (p.get("text", "#111"), p.get("text2", "#555"),
                p.get("border", "#ddd"), p.get("primary", ZIELEN_GLOWNA))
    except Exception:
        return "#111", "#555", "#ddd", ZIELEN_GLOWNA


def _fmt(iso: str) -> str:
    try:
        return dt.date.fromisoformat(str(iso)[:10]).strftime("%d.%m.%Y")
    except Exception:
        return str(iso)


def _pasek_sortowania(prefix: str, kolumny: list[str], domyslna: str) -> tuple[str, bool]:
    """Klikalne nagłówki kolumn zamiast list wyboru. Klik ustawia sortowanie po
    danej kolumnie (malejąco); ponowny klik w tę samą kolumnę odwraca kierunek.
    Aktywna kolumna oznaczona strzałką ▼/▲. Zwraca (kolumna, malejąco)."""
    kl_kol, kl_dir = f"{prefix}_sortkol", f"{prefix}_malej"
    if kl_kol not in st.session_state:
        st.session_state[kl_kol] = domyslna
    if kl_dir not in st.session_state:
        st.session_state[kl_dir] = True

    st.caption("Sortuj — kliknij kolumnę (ponowny klik odwraca kolejność):")
    kolumny_widget = st.columns(len(kolumny))
    for c, etykieta in zip(kolumny_widget, kolumny):
        aktywna = st.session_state[kl_kol] == etykieta
        strzalka = (" ▼" if st.session_state[kl_dir] else " ▲") if aktywna else ""
        if c.button(etykieta + strzalka, key=f"{prefix}_sb_{etykieta}",
                    use_container_width=True,
                    type="primary" if aktywna else "secondary"):
            if aktywna:
                st.session_state[kl_dir] = not st.session_state[kl_dir]
            else:
                st.session_state[kl_kol] = etykieta
                st.session_state[kl_dir] = True
            st.rerun()
    return st.session_state[kl_kol], st.session_state[kl_dir]


def _tabela_html(rekordy: list[dict]) -> str:
    txt, txt2, bord, head = _kolory()
    th = (f"padding:8px 10px;border-bottom:2px solid {head};color:{head};"
          f"text-align:left;font-size:0.85rem;font-weight:700;")
    style = f"width:100%;border-collapse:collapse;font-size:0.9rem;color:{txt};"

    # Kolumny warunkowe — pojawiają się tylko, gdy jakikolwiek wiersz je ma.
    z_branza = any((r.get("branza") or "").strip() for r in rekordy)
    z_przedmiot = any((r.get("przedmiot") or "").strip() for r in rekordy)
    # „Data publikacji” (= data pobrania do bazy) — pokazywana w module 6;
    # w module 5 rekordy jej nie mają, więc kolumna się nie renderuje.
    z_publikacja = any(r.get("data_publikacji") for r in rekordy)

    kol_branza_naglowek = (
        f"<th style='{th}white-space:nowrap;'>Branża</th>" if z_branza else "")
    kol_przedmiot_naglowek = (
        f"<th style='{th}'>Przedmiot</th>" if z_przedmiot else "")
    kol_publikacja_naglowek = (
        f"<th style='{th}white-space:nowrap;'>Data publikacji</th>"
        if z_publikacja else "")
    naglowek = (
        f"<tr>"
        f"<th style='{th}width:36px;'>L.p.</th>"
        f"<th style='{th}white-space:nowrap;'>Sygnatura</th>"
        f"<th style='{th}white-space:nowrap;'>Data wydania</th>"
        f"{kol_publikacja_naglowek}"
        f"<th style='{th}'>Temat</th>"
        f"{kol_branza_naglowek}"
        f"{kol_przedmiot_naglowek}"
        f"<th style='{th}'>Streszczenie</th>"
        f"</tr>"
    )

    wiersze = []
    for i, r in enumerate(rekordy, start=1):
        td = (f"padding:8px 10px;border-bottom:1px solid {bord};"
              f"vertical-align:top;color:{txt};")

        kom_publikacja = ""
        if z_publikacja:
            pub = _fmt(r["data_publikacji"]) if r.get("data_publikacji") else "—"
            rozna = (r.get("data_publikacji") and
                     str(r.get("data_publikacji"))[:10] != str(r.get("data_wyd"))[:10])
            if rozna:
                pub = (f"{pub}<br><span style='font-size:0.72rem;color:{txt2};'>"
                       f"inna niż data wydania</span>")
            kom_publikacja = f"<td style='{td}white-space:nowrap;'>{pub}</td>"

        kom_branza = ""
        if z_branza:
            br = html.escape((r.get("branza") or "").strip())
            kom_branza = (f"<td style='{td}white-space:nowrap;"
                          f"font-size:0.82rem;'>{br}</td>")
        kom_przedmiot = ""
        if z_przedmiot:
            pr = html.escape((r.get("przedmiot") or "").strip())
            kom_przedmiot = f"<td style='{td}font-size:0.82rem;'>{pr}</td>"

        wiersze.append(
            f"<tr>"
            f"<td style='{td}'>{i}</td>"
            f"<td style='{td}white-space:nowrap;font-family:monospace;font-size:0.82rem;'>"
            f"{html.escape(r['sygnatura'])}</td>"
            f"<td style='{td}white-space:nowrap;'>{_fmt(r['data_wyd'])}</td>"
            f"{kom_publikacja}"
            f"<td style='{td}'>{html.escape(r.get('temat') or '')}</td>"
            f"{kom_branza}"
            f"{kom_przedmiot}"
            f"<td style='{td}'>{html.escape(r.get('streszczenie') or '')}</td>"
            f"</tr>"
        )

    return f"<table style='{style}'>{naglowek}{''.join(wiersze)}</table>"


# ---------------------------------------------------------------------------
# RENDER JEDNEJ ZAKŁADKI
# ---------------------------------------------------------------------------
def _moze_wgrywac() -> bool:
    """Wgrywanie plików to uprawnienie administratora. Zwykły użytkownik
    widzi zestawienia (tabele), ale nie wgrywa."""
    rola = "admin" if (st.session_state.get("superadmin")
                       or st.session_state.get("rola") == "admin") else "user"
    return auth.ma_uprawnienie(rola, "zestawienie_wgrywanie")


def _zakladka(podatek: str) -> None:
    # ── WGRYWANIE (tylko administrator; user od razu przechodzi do tabel) ────
    if _moze_wgrywac():
        plik = st.file_uploader(
            f"Wgraj plik „Tygodniowego Researchu” dla {podatek} (DOCX zalecany, PDF też)",
            type=["docx", "pdf"],
            key=f"up_{podatek}",
            help="Moduł sam odczyta daty wydania i przypisze interpretacje do właściwych tygodni.",
        )
        if plik is not None:
            bajty = plik.getvalue()
            h = hashlib.md5(bajty).hexdigest()
            if st.session_state.get(f"hash_{podatek}") != h:
                with st.spinner("Parsuję i przypisuję do tygodni…"):
                    wiersze = _parsuj_plik(plik.name, bajty)
                    if not wiersze:
                        st.error(
                            "Nie udało się odczytać tabeli z pliku. Upewnij się, że to "
                            "plik z tabelą (Sygnatura / Data wydania / Temat / Streszczenie). "
                            "Najpewniej działa DOCX."
                        )
                    else:
                        zap, pom, tyg = _zapisz_interpretacje(podatek, plik.name, wiersze)
                        st.session_state[f"hash_{podatek}"] = h
                        kom = f"Przypisano {zap} interpretacji"
                        if tyg:
                            kom += f" (okres researchu: {_etykieta_tygodnia(tyg)})"
                        if pom:
                            kom += f"; pominięto {pom} bez czytelnej daty wydania"
                        st.success(kom + ".")
                        st.rerun()
        st.divider()

    # ── SORTOWANA TABELA (klik w nagłówek kolumny) ──────────────────────────
    sort_kol, malejaco = _pasek_sortowania(
        f"tyg_{podatek}", list(SORT_KOLUMNY_5.keys()), "Data wydania")
    rekordy = _interpretacje_sortowane(podatek, sort_kol, malejaco)
    if not rekordy:
        st.info("Brak wgranych streszczeń dla tego podatku. Wgraj plik powyżej.")
        return

    st.markdown(_tabela_html(rekordy), unsafe_allow_html=True)
    st.caption(f"Pokazano {len(rekordy)} interpretacji {podatek} "
               f"(limit {LIMIT_WIERSZY_5}, sortowanie: {sort_kol.lower()}).")


# ---------------------------------------------------------------------------
# PUNKT WEJŚCIA
# ---------------------------------------------------------------------------
def pokaz_zestawienie_tygodniowe() -> None:
    st.header("📅 Zestawienie tygodniowe")
    st.caption(
        "Wgraj plik z GPT „Tygodniowy Research”. Interpretacje trafią do "
        "wspólnej tabeli, którą sortujesz po dacie wydania lub sygnaturze."
    )

    try:
        _zapewnij_tabele()
    except Exception as e:
        st.error(f"Nie udało się przygotować tabeli: {e}")
        return

    for zakladka_ui, podatek in zip(st.tabs(PODATKI), PODATKI):
        with zakladka_ui:
            _zakladka(podatek)


if __name__ == "__main__":
    st.set_page_config(page_title="Zestawienie tygodniowe", layout="wide")
    pokaz_zestawienie_tygodniowe()
