"""
raport_silnik.py — Wspolna logika generowania raportu dla DOWOLNEGO zakresu dat i podatku.

Uzywany przez:
  - raport_tygodniowy.py    (cron, caly tydzien, wszystkie podatki)
  - raport_na_zadanie.py    (GitHub Actions workflow_dispatch, jeden rok/miesiac/podatek)
  - raporty.py               (Streamlit, generowanie "na zywo" w aplikacji)

Nie zalezy od Streamlit ani od GitHub Actions - czysta logika biznesowa.
"""

import io
import os
import time
import smtplib
import requests
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from docx import Document

import db_core
import utils


PODATKI_WSZYSTKIE = ["PIT", "CIT", "VAT", "AKCYZA"]


# ---------------------------------------------------------------------------
# GENEROWANIE WORD
# ---------------------------------------------------------------------------
def generuj_word(rekordy: list, podatek: str, opis_okresu: str, tytul_raportu: str = "Raport") -> bytes:
    """
    rekordy: lista slownikow z kluczami Sygnatura/Data/Podatek/Link/Tekst/Format
    opis_okresu: czytelny string np. "Styczen 2026" lub "06.01 — 10.01.2026"
    """
    doc = Document()
    doc.add_heading(f"PickPivot — {tytul_raportu}: {podatek}", 0)
    doc.add_paragraph(f"Okres: {opis_okresu}")
    doc.add_paragraph(f"Wygenerowano: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Liczba dokumentow: {len(rekordy)}")
    doc.add_page_break()

    for r in sorted(rekordy, key=lambda x: x.get("Data", ""), reverse=True):
        doc.add_heading(f"Sygnatura: {r['Sygnatura']}", 1)
        doc.add_paragraph(f"Data:    {r['Data']}")
        doc.add_paragraph(f"Podatek: {r['Podatek']}")
        doc.add_paragraph(f"Link:    {r['Link']}")
        if r.get("Format"):
            doc.add_paragraph(f"Zrodlo:  {r['Format']}")
        doc.add_paragraph(utils.wyczysc_tekst_dla_worda(r["Tekst"]))
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# UZUPELNIANIE ARCHIWUM Z API MF DLA DOWOLNEGO ZAKRESU DAT
# ---------------------------------------------------------------------------
def uzupelnij_archiwum(
    db: db_core.SupabaseDB,
    podatek: str,
    data_od: datetime,
    data_do: datetime,
    log_fn=print,
    workers: int = 4,
) -> tuple:
    """
    Sprawdza API MF dla okresu i pobiera brakujace dokumenty do archiwum.
    Zwraca (liczba_nowych, status: "OK"|"NIEPELNE_POBRANIE"|"BLOKADA"|"ERROR").
    """
    log_fn(f"[{podatek}] Sprawdzam API MF dla okresu {data_od.date()} — {data_do.date()}...")

    znane_id = db_core.pobierz_id_z_archiwum(db)

    with requests.Session() as sesja:
        lista, status = utils.pobierz_wszystko_z_okresu(
            data_od.strftime("%Y-%m-%d"),
            data_do.strftime("%Y-%m-%d"),
            sesja, podatek, utils.KODY_PRZEPISOW[podatek],
        )

    if status != "OK":
        log_fn(f"[{podatek}] OSTRZEZENIE: API MF zwrocilo status {status}")
        return 0, status

    do_pobrania = [d for d in lista if d["id"] not in znane_id]
    log_fn(f"[{podatek}] Znaleziono {len(lista)} w MF, do pobrania: {len(do_pobrania)}")

    if not do_pobrania:
        log_fn(f"[{podatek}] Archiwum juz aktualne dla tego okresu.")
        return 0, "OK"

    def on_postep(completed, total, sygnatura, status):
        if completed % 10 == 0 or completed == total:
            log_fn(f"[{podatek}] Postep: {completed}/{total}")

    nowe_tresci, _, _, blokada = utils.pobierz_dokumenty_rownolegle(
        do_pobrania, znane_id, set(),
        callback_postep=on_postep, workers=workers,
    )

    if nowe_tresci:
        zapisanych = db_core.zapisz_wiele_do_archiwum(db, nowe_tresci, "raport_na_zadanie")
        log_fn(f"[{podatek}] Zapisano {zapisanych} nowych dokumentow.")
        return zapisanych, ("BLOKADA" if blokada else "OK")

    return 0, ("BLOKADA" if blokada else "OK")


# ---------------------------------------------------------------------------
# WERYFIKACJA KOMPLETNOSCI — drugie, niezalezne zapytanie do API MF
# ---------------------------------------------------------------------------
def weryfikuj_kompletnosc(
    podatek: str,
    data_od: datetime,
    data_do: datetime,
    liczba_w_archiwum: int,
    log_fn=print,
    max_prob: int = 3,
) -> dict:
    """
    Odpytuje API MF JESZCZE RAZ (niezaleznie od pobierania) tylko po to,
    zeby porownac ile dokumentow MF zglasza dla tego okresu z tym co mamy
    w archiwum. To wykrywa sytuacje, w ktorych pierwsze zapytanie zwrocilo
    niepelna liste (np. przez cichy blad MF bez kodu bledu HTTP).

    Zwraca dict:
        {
            "zgodnosc": bool,           # True jesli liczby sie zgadzaja
            "liczba_w_mf": int | None,  # None jesli weryfikacja sie nie powiodla
            "liczba_w_archiwum": int,
            "roznica": int,
            "status": "OK" | "NIEZGODNOSC" | "WERYFIKACJA_NIEUDANA"
        }
    """
    log_fn(f"[{podatek}] Weryfikacja kompletnosci — drugie zapytanie do API MF...")

    for proba in range(1, max_prob + 1):
        try:
            with requests.Session() as sesja:
                lista, status = utils.pobierz_wszystko_z_okresu(
                    data_od.strftime("%Y-%m-%d"),
                    data_do.strftime("%Y-%m-%d"),
                    sesja, podatek, utils.KODY_PRZEPISOW[podatek],
                )

            # "OK" i "NIEPELNE_POBRANIE" oba dostarczaja faktyczna liste -
            # NIEPELNE_POBRANIE oznacza ze paginacja nie zebrala wszystkiego
            # co API zadeklarowalo w totalHits, wiec liczba_w_mf moze byc
            # zanizona. Traktujemy to jako sygnal do ponowienia, nie jako
            # wiarygodny wynik.
            if status == "OK":
                liczba_w_mf = len(lista)
                roznica = liczba_w_mf - liczba_w_archiwum

                if roznica == 0:
                    log_fn(f"[{podatek}] Weryfikacja OK — MF: {liczba_w_mf}, archiwum: {liczba_w_archiwum}.")
                    return {
                        "zgodnosc": True, "liczba_w_mf": liczba_w_mf,
                        "liczba_w_archiwum": liczba_w_archiwum, "roznica": 0, "status": "OK",
                    }
                else:
                    log_fn(
                        f"[{podatek}] NIEZGODNOSC — MF zglasza {liczba_w_mf}, "
                        f"w archiwum {liczba_w_archiwum} (roznica: {roznica})."
                    )
                    return {
                        "zgodnosc": False, "liczba_w_mf": liczba_w_mf,
                        "liczba_w_archiwum": liczba_w_archiwum, "roznica": roznica,
                        "status": "NIEZGODNOSC",
                    }

            if status == "NIEPELNE_POBRANIE":
                log_fn(
                    f"[{podatek}] Proba weryfikacji {proba}/{max_prob} — paginacja API "
                    f"nie zebrala wszystkich zadeklarowanych wynikow, ponawiam..."
                )
            else:
                log_fn(f"[{podatek}] Proba weryfikacji {proba}/{max_prob} — status {status}, ponawiam...")
            time.sleep(8 * proba)

        except Exception as e:
            log_fn(f"[{podatek}] Blad podczas weryfikacji (proba {proba}): {e}")
            time.sleep(8 * proba)

    log_fn(f"[{podatek}] Weryfikacja nieudana po {max_prob} probach — MF niedostepne.")
    return {
        "zgodnosc": None, "liczba_w_mf": None,
        "liczba_w_archiwum": liczba_w_archiwum, "roznica": None,
        "status": "WERYFIKACJA_NIEUDANA",
    }


# ---------------------------------------------------------------------------
# GENEROWANIE RAPORTU DLA JEDNEGO PODATKU + ZAKRESU DAT (rdzen logiki)
# ---------------------------------------------------------------------------
def generuj_raport_dla_podatku(
    db: db_core.SupabaseDB,
    podatek: str,
    data_od: datetime,
    data_do: datetime,
    opis_okresu: str,
    log_fn=print,
    weryfikuj: bool = True,
    generuj_plik: bool = True,
) -> dict:
    """
    Pelny przebieg dla JEDNEGO podatku: uzupelnia archiwum, weryfikuje
    kompletnosc (drugie zapytanie do MF), opcjonalnie generuje Word, zwraca wynik.

    generuj_plik: jesli False, pomija budowanie pliku Word (plik_bytes zawsze
    None) - uzywane przez "Sciagacz Interpretacji" -> Pobieranie na zadanie,
    ktore ma tylko wgrywac dane do bazy i wysylac powiadomienie mailem, bez
    generowania dokumentu. raport_tygodniowy.py nie przekazuje tego parametru,
    wiec dziala dokladnie tak jak wczesniej (generuj_plik=True domyslnie).

    Zwraca dict:
        {
            "podatek": str, "liczba_dok": int, "plik_bytes": bytes|None,
            "nowych_pobranych": int, "status": "OK"|"BLOKADA"|"ERROR"|"BRAK_DOKUMENTOW",
            "weryfikacja": dict | None  # wynik weryfikuj_kompletnosc()
        }
    """
    try:
        nowych, status_pobierania = uzupelnij_archiwum(db, podatek, data_od, data_do, log_fn=log_fn)

        rekordy = db_core.pobierz_rekordy_z_archiwum(
            db, podatek=podatek,
            data_od=data_od.strftime("%Y-%m-%d"),
            data_do=data_do.strftime("%Y-%m-%d"),
        )

        if not rekordy:
            return {
                "podatek": podatek, "liczba_dok": 0, "plik_bytes": None,
                "nowych_pobranych": nowych, "status": "BRAK_DOKUMENTOW",
                "weryfikacja": None,
            }

        # ── PODWOJNA WERYFIKACJA — niezalezne drugie zapytanie do MF ──────
        wynik_weryfikacji = None
        if weryfikuj:
            wynik_weryfikacji = weryfikuj_kompletnosc(
                podatek, data_od, data_do, len(rekordy), log_fn=log_fn
            )

            # Jesli wykryto niezgodnosc - sprobuj DOCIAGNAC brakujace dokumenty
            # zanim wygenerujemy finalny plik (samoleczace zachowanie)
            if wynik_weryfikacji["status"] == "NIEZGODNOSC" and wynik_weryfikacji["roznica"] > 0:
                log_fn(f"[{podatek}] Wykryto brakujace dokumenty — proba douzupelnienia archiwum...")
                nowych_dodatkowo, _ = uzupelnij_archiwum(db, podatek, data_od, data_do, log_fn=log_fn)
                nowych += nowych_dodatkowo

                if nowych_dodatkowo > 0:
                    rekordy = db_core.pobierz_rekordy_z_archiwum(
                        db, podatek=podatek,
                        data_od=data_od.strftime("%Y-%m-%d"),
                        data_do=data_do.strftime("%Y-%m-%d"),
                    )
                    # Ponowna weryfikacja po douzupelnieniu
                    wynik_weryfikacji = weryfikuj_kompletnosc(
                        podatek, data_od, data_do, len(rekordy), log_fn=log_fn, max_prob=2
                    )

        plik_bytes = None
        if generuj_plik:
            plik_bytes = generuj_word(rekordy, podatek, opis_okresu, tytul_raportu="Raport na zadanie")

        # Status koncowy uwzglednia wynik weryfikacji
        status_finalny = status_pobierania
        if wynik_weryfikacji and wynik_weryfikacji["status"] == "NIEZGODNOSC":
            status_finalny = "NIEZGODNOSC"
        elif wynik_weryfikacji and wynik_weryfikacji["status"] == "WERYFIKACJA_NIEUDANA":
            status_finalny = "WERYFIKACJA_NIEUDANA"

        return {
            "podatek": podatek, "liczba_dok": len(rekordy), "plik_bytes": plik_bytes,
            "nowych_pobranych": nowych, "status": status_finalny,
            "weryfikacja": wynik_weryfikacji,
        }

    except Exception as e:
        log_fn(f"[{podatek}] BLAD: {e}")
        return {
            "podatek": podatek, "liczba_dok": 0, "plik_bytes": None,
            "nowych_pobranych": 0, "status": "ERROR", "error_msg": str(e),
            "weryfikacja": None,
        }


# ---------------------------------------------------------------------------
# WYSYLKA POWIADOMIENIA MAILEM — BEZ ZALACZNIKA (Sciagacz Interpretacji)
# ---------------------------------------------------------------------------
def wyslij_email_powiadomienie_pobrania(
    wyniki: list,
    opis_okresu: str,
    gmail_adres: str,
    gmail_haslo: str,
    odbiorca: str,
    log_fn=print,
) -> bool:
    """
    Wysyla PROSTE powiadomienie e-mail o zakonczonym pobraniu interpretacji
    do bazy danych — BEZ zalacznika Word. Uzywane przez "Sciagacz
    Interpretacji" -> Pobieranie na zadanie, gdzie celem jest wylacznie
    zasilenie bazy danymi, nie wygenerowanie dokumentu.

    wyniki: lista dict z generuj_raport_dla_podatku() (wywolanej z generuj_plik=False)
    Zwraca True przy udanej wysylce.
    """
    podatki_str = "/".join(w["podatek"] for w in wyniki)
    ma_niezgodnosc = any(
        w.get("weryfikacja") and w["weryfikacja"]["status"] in ("NIEZGODNOSC", "WERYFIKACJA_NIEUDANA")
        for w in wyniki
    )
    temat_status = " [wymaga sprawdzenia]" if ma_niezgodnosc else ""
    temat = f"PickPivot — Pobrano do bazy: {podatki_str} ({opis_okresu}){temat_status}"

    def _status_kom(w):
        wer = w.get("weryfikacja")
        if not wer:
            return "brak weryfikacji"
        mapa = {
            "OK": "✅ potwierdzona kompletność",
            "NIEZGODNOSC": f"⚠️ rozbieżność (różnica: {wer['roznica']})",
            "WERYFIKACJA_NIEUDANA": "ℹ️ niepotwierdzone (MF niedostępne)",
        }
        return mapa.get(wer["status"], wer["status"])

    wiersze = "".join(
        f'<tr><td>{w["podatek"]}</td><td>{w["liczba_dok"]}</td>'
        f'<td>{w.get("nowych_pobranych", 0)}</td>'
        f'<td>{_status_kom(w)}</td></tr>'
        for w in wyniki
    )

    tresc_html = f"""
    <html><body style="font-family: Arial, sans-serif;">
    <h2>📥 PickPivot — Ściągacz Interpretacji: pobieranie na żądanie</h2>
    <p><b>Okres:</b> {opis_okresu}</p>
    <p>Interpretacje zostały pobrane i zapisane w bazie danych. Podsumowanie:</p>
    <table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
        <tr style="background-color: #2C3E50; color: white;">
            <th>Podatek</th><th>Łącznie w bazie (ten okres)</th><th>Nowo pobranych</th><th>Weryfikacja</th>
        </tr>
        {wiersze}
    </table>
    <p style="color: #888; font-size: 12px; margin-top: 20px;">
        Wiadomość wygenerowana automatycznie przez PickPivot — Ściągacz Interpretacji.
    </p>
    </body></html>
    """

    msg = MIMEMultipart("alternative")
    msg["Subject"] = temat
    msg["From"]    = gmail_adres
    msg["To"]      = odbiorca
    msg.attach(MIMEText(tresc_html, "html", "utf-8"))

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_adres, gmail_haslo)
            server.send_message(msg)
        log_fn(f"E-mail powiadamiający wysłany do {odbiorca}.")
        return True
    except Exception as e:
        log_fn(f"BŁĄD wysyłki email: {e}")
        return False


# ---------------------------------------------------------------------------
# WYSYLKA POWIADOMIENIA — CODZIENNA SYNCHRONIZACJA AUTOMATYCZNA
# ---------------------------------------------------------------------------
def wyslij_email_synchronizacja_dzienna(
    wyniki: list,
    opis_okresu: str,
    gmail_adres: str,
    gmail_haslo: str,
    odbiorca: str,
    log_fn=print,
) -> bool:
    """
    Krotkie, codzienne powiadomienie o wyniku automatycznej synchronizacji
    nocnej (3:00). Wizualnie odrebne od powiadomien "na zadanie", zeby
    odbiorca od razu widzial w skrzynce ktore maile sa automatyczne.

    wyniki: lista dict z generuj_raport_dla_podatku() (generuj_plik=False)
    """
    nowych_lacznie = sum(w.get("nowych_pobranych", 0) for w in wyniki)
    ma_niezgodnosc = any(
        w.get("weryfikacja") and w["weryfikacja"]["status"] in ("NIEZGODNOSC", "WERYFIKACJA_NIEUDANA")
        for w in wyniki
    )
    ma_blad = any(w["status"] == "ERROR" for w in wyniki)

    if ma_blad:
        znacznik = "❌"
    elif ma_niezgodnosc:
        znacznik = "⚠️"
    elif nowych_lacznie > 0:
        znacznik = "🆕"
    else:
        znacznik = "✅"

    temat = f"{znacznik} PickPivot — Synchronizacja dzienna {opis_okresu} ({nowych_lacznie} nowych)"

    def _status_kom(w):
        wer = w.get("weryfikacja")
        if w["status"] == "ERROR":
            return "❌ błąd"
        if not wer:
            return "—"
        mapa = {
            "OK": "✅ potwierdzona kompletność",
            "NIEZGODNOSC": f"⚠️ rozbieżność (różnica: {wer['roznica']})",
            "WERYFIKACJA_NIEUDANA": "ℹ️ niepotwierdzone (MF niedostępne)",
        }
        return mapa.get(wer["status"], wer["status"])

    wiersze = "".join(
        f'<tr><td>{w["podatek"]}</td><td>{w["liczba_dok"]}</td>'
        f'<td>{w.get("nowych_pobranych", 0)}</td>'
        f'<td>{_status_kom(w)}</td></tr>'
        for w in wyniki
    )

    tresc_html = f"""
    <html><body style="font-family: Arial, sans-serif;">
    <h2>🌙 PickPivot — Synchronizacja dzienna (automatyczna, 3:00)</h2>
    <p><b>Sprawdzony okres:</b> {opis_okresu} (ruchome okno 3 dni)</p>
    <table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
        <tr style="background-color: #2C3E50; color: white;">
            <th>Podatek</th><th>Łącznie w bazie (ten okres)</th><th>Nowo pobranych dziś</th><th>Weryfikacja</th>
        </tr>
        {wiersze}
    </table>
    <p style="color: #888; font-size: 12px; margin-top: 20px;">
        Wiadomość wygenerowana automatycznie o 3:00 przez PickPivot — Ściągacz Interpretacji
        (codzienna synchronizacja, GitHub Actions).
    </p>
    </body></html>
    """

    msg = MIMEMultipart("alternative")
    msg["Subject"] = temat
    msg["From"]    = gmail_adres
    msg["To"]      = odbiorca
    msg.attach(MIMEText(tresc_html, "html", "utf-8"))

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_adres, gmail_haslo)
            server.send_message(msg)
        log_fn(f"E-mail synchronizacji dziennej wysłany do {odbiorca}.")
        return True
    except Exception as e:
        log_fn(f"BŁĄD wysyłki email: {e}")
        return False


# ---------------------------------------------------------------------------
# WYSYLKA EMAIL Z ZALACZNIKIEM
# ---------------------------------------------------------------------------
def wyslij_email_z_zalacznikiem(
    plik_bytes: bytes,
    nazwa_pliku: str,
    podatek: str,
    opis_okresu: str,
    liczba_dok: int,
    gmail_adres: str,
    gmail_haslo: str,
    odbiorca: str,
    log_fn=print,
    weryfikacja: dict = None,
) -> bool:
    """Wysyla e-mail z plikiem Word jako zalacznik. Zwraca True przy sukcesie."""
    from email.mime.application import MIMEApplication

    badge_weryfikacji = _html_badge_weryfikacji(weryfikacja)
    temat_status = ""
    if weryfikacja:
        if weryfikacja["status"] == "NIEZGODNOSC":
            temat_status = " [⚠️ WYMAGA SPRAWDZENIA]"
        elif weryfikacja["status"] == "WERYFIKACJA_NIEUDANA":
            temat_status = " [ℹ️ niepotwierdzone]"

    temat = f"PickPivot — Raport na zadanie: {podatek} ({opis_okresu}){temat_status}"

    tresc_html = f"""
    <html><body style="font-family: Arial, sans-serif;">
    <h2>📄 PickPivot — Raport na żądanie</h2>
    <p><b>Podatek:</b> {podatek}</p>
    <p><b>Okres:</b> {opis_okresu}</p>
    <p><b>Liczba dokumentów:</b> {liczba_dok}</p>
    {badge_weryfikacji}
    <p style="margin-top: 20px;">Plik Word w załączniku.</p>
    <p style="color: #888; font-size: 12px;">
        Wiadomość wygenerowana na żądanie użytkownika PickPivot.
    </p>
    </body></html>
    """

    msg = MIMEMultipart("mixed")
    msg["Subject"] = temat
    msg["From"]    = gmail_adres
    msg["To"]      = odbiorca

    alt = MIMEMultipart("alternative")
    alt.attach(MIMEText(tresc_html, "html", "utf-8"))
    msg.attach(alt)

    zalacznik = MIMEApplication(plik_bytes, _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document")
    zalacznik.add_header("Content-Disposition", "attachment", filename=nazwa_pliku)
    msg.attach(zalacznik)

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_adres, gmail_haslo)
            server.send_message(msg)
        log_fn(f"E-mail z raportem ({podatek}) wyslany do {odbiorca}.")
        return True
    except Exception as e:
        log_fn(f"BLAD wysylki email: {e}")
        return False


def _html_badge_weryfikacji(weryfikacja: dict) -> str:
    """Generuje kolorowy badge HTML pokazujacy status weryfikacji kompletnosci."""
    if not weryfikacja:
        return ""

    status = weryfikacja["status"]

    if status == "OK":
        return (
            '<p style="background:#D5F5E3; padding:10px; border-radius:6px; border-left:4px solid #27AE60;">'
            '✅ <b>Potwierdzona kompletność</b> — niezależna druga weryfikacja w API MF '
            f'zgadza się z liczbą dokumentów w raporcie ({weryfikacja["liczba_w_mf"]}).</p>'
        )
    elif status == "NIEZGODNOSC":
        znak = "+" if weryfikacja["roznica"] > 0 else ""
        return (
            '<p style="background:#FADBD8; padding:10px; border-radius:6px; border-left:4px solid #E74C3C;">'
            '⚠️ <b>Wymaga sprawdzenia</b> — druga weryfikacja w API MF wykryła rozbieżność: '
            f'MF zgłasza {weryfikacja["liczba_w_mf"]} dokumentów, w raporcie jest '
            f'{weryfikacja["liczba_w_archiwum"]} (różnica: {znak}{weryfikacja["roznica"]}). '
            'System próbował automatycznie douzupełnić archiwum — sprawdź czy liczby się zgadzają, '
            'a w razie wątpliwości uruchom raport ponownie.</p>'
        )
    elif status == "WERYFIKACJA_NIEUDANA":
        return (
            '<p style="background:#FEF9E7; padding:10px; border-radius:6px; border-left:4px solid #F39C12;">'
            'ℹ️ <b>Kompletność niepotwierdzona</b> — serwer Ministerstwa Finansów był niedostępny '
            'podczas próby weryfikacji. Raport zawiera wszystko co udało się pobrać przy pierwszym '
            'zapytaniu, ale zalecamy ponowne uruchomienie raportu za jakiś czas dla pewności.</p>'
        )
    return ""


def wyslij_email_podsumowanie_wielu(
    wyniki: list,
    opis_okresu: str,
    gmail_adres: str,
    gmail_haslo: str,
    odbiorca: str,
    log_fn=print,
) -> bool:
    """
    Wysyla JEDEN email podsumowujacy gdy generowano kilka podatkow naraz,
    z wszystkimi plikami jako zalaczniki.
    wyniki: lista dict z generuj_raport_dla_podatku()
    """
    from email.mime.application import MIMEApplication

    podatki_str = "/".join(w["podatek"] for w in wyniki)
    ma_niezgodnosc = any(
        w.get("weryfikacja") and w["weryfikacja"]["status"] in ("NIEZGODNOSC", "WERYFIKACJA_NIEUDANA")
        for w in wyniki
    )
    temat_status = " [⚠️ SPRAWDZ SZCZEGOLY]" if ma_niezgodnosc else ""
    temat = f"PickPivot — Raport na zadanie: {podatki_str} ({opis_okresu}){temat_status}"

    def _status_kom(w):
        wer = w.get("weryfikacja")
        if not wer:
            return "—"
        if wer["status"] == "OK":
            return "✅ Potwierdzone"
        if wer["status"] == "NIEZGODNOSC":
            znak = "+" if wer["roznica"] > 0 else ""
            return f"⚠️ Różnica {znak}{wer['roznica']}"
        if wer["status"] == "WERYFIKACJA_NIEUDANA":
            return "ℹ️ Niepotwierdzone"
        return "—"

    wiersze = "".join(
        f'<tr><td>{w["podatek"]}</td><td>{w["liczba_dok"]}</td>'
        f'<td>{"✅ Zalaczony" if w["plik_bytes"] else "— brak dokumentow"}</td>'
        f'<td>{_status_kom(w)}</td></tr>'
        for w in wyniki
    )

    info_dolna = (
        '<p style="background:#FADBD8; padding:10px; border-radius:6px; border-left:4px solid #E74C3C; margin-top:15px;">'
        '⚠️ Co najmniej jeden podatek ma niepotwierdzoną kompletność lub wykrytą rozbieżność — '
        'sprawdź kolumnę "Weryfikacja" w tabeli powyżej.</p>'
        if ma_niezgodnosc else ""
    )

    tresc_html = f"""
    <html><body style="font-family: Arial, sans-serif;">
    <h2>📄 PickPivot — Raport na żądanie</h2>
    <p><b>Okres:</b> {opis_okresu}</p>
    <table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse;">
        <tr style="background-color: #2C3E50; color: white;">
            <th>Podatek</th><th>Liczba dokumentów</th><th>Plik</th><th>Weryfikacja</th>
        </tr>
        {wiersze}
    </table>
    {info_dolna}
    <p style="margin-top: 20px;">Pliki Word w załączniku (jeśli były dokumenty).</p>
    <p style="color: #888; font-size: 12px;">
        Wiadomość wygenerowana na żądanie użytkownika PickPivot.
    </p>
    </body></html>
    """

    msg = MIMEMultipart("mixed")
    msg["Subject"] = temat
    msg["From"]    = gmail_adres
    msg["To"]      = odbiorca

    alt = MIMEMultipart("alternative")
    alt.attach(MIMEText(tresc_html, "html", "utf-8"))
    msg.attach(alt)

    for w in wyniki:
        if w["plik_bytes"]:
            nazwa = f"Raport_{w['podatek']}_{opis_okresu.replace(' ', '_').replace('—','-')}.docx"
            zalacznik = MIMEApplication(
                w["plik_bytes"],
                _subtype="vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            zalacznik.add_header("Content-Disposition", "attachment", filename=nazwa)
            msg.attach(zalacznik)

    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(gmail_adres, gmail_haslo)
            server.send_message(msg)
        log_fn(f"E-mail podsumowujacy wyslany do {odbiorca}.")
        return True
    except Exception as e:
        log_fn(f"BLAD wysylki email: {e}")
        return False


# ---------------------------------------------------------------------------
# POMOCNICZE — zakres dat z roku/miesiaca
# ---------------------------------------------------------------------------
def zakres_z_roku_miesiaca(rok: int, miesiac: int) -> tuple:
    """Zwraca (data_od, data_do, opis_okresu) dla calego miesiaca."""
    import calendar as cal_mod
    data_od = datetime(rok, miesiac, 1)
    ostatni_dzien = cal_mod.monthrange(rok, miesiac)[1]
    data_do = datetime(rok, miesiac, ostatni_dzien)

    miesiace_pl = ["Styczen","Luty","Marzec","Kwiecien","Maj","Czerwiec",
                   "Lipiec","Sierpien","Wrzesien","Pazdziernik","Listopad","Grudzien"]
    opis = f"{miesiace_pl[miesiac-1]} {rok}"
    return data_od, data_do, opis


def zakres_3_dni() -> tuple:
    """
    Zwraca (data_od, data_do, opis_okresu) dla ruchomego okna 3 dni
    konczacego sie dzisiaj (dzisiaj, wczoraj, przedwczoraj).

    Uzywane przez codzienna synchronizacje - kazdy dzien jest sprawdzany
    3-krotnie w kolejnych uruchomieniach (jako "dzisiaj", potem "wczoraj",
    potem "przedwczoraj"), co daje odpornosc na interpretacje publikowane
    przez MF z data wsteczna. Duplikaty sa pomijane automatycznie przy
    zapisie (ON CONFLICT DO NOTHING), wiec powtorne sprawdzanie tych samych
    dni jest bezpieczne i nie generuje zduplikowanych rekordow.
    """
    from datetime import timedelta
    dzisiaj = datetime.now()
    data_od = dzisiaj - timedelta(days=2)
    data_do = dzisiaj
    opis = f"{data_od.strftime('%d.%m')} — {data_do.strftime('%d.%m.%Y')}"
    return data_od, data_do, opis
